"""
Async Practice API for SkillBot
Practice session management endpoints with analytics
"""

import logging
import aiofiles
from typing import Optional, Dict, Any, List
from datetime import datetime, timedelta
from bson import ObjectId

from fastapi import APIRouter, Request, HTTPException, Depends, status, Query
from pydantic import BaseModel, Field, validator, root_validator
from slowapi import Limiter
from slowapi.util import get_remote_address

from core.database import DatabaseManager
from core.cache import CacheManager
from api.v1.auth_async import get_current_user, get_database, get_cache
from config_async import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# LLM Gate bridge (SWM-011)
#
# All LLM calls — text-only AND vision/multimodal — MUST be routed through
# the shared gate (C4) with caller_id ``pcr_practice``.  If the gate module
# is available, it is the exclusive path — there is NO silent fallback to a
# direct provider.
#
# Vision/multimodal calls use ``gate.call(messages=...)`` which forwards the
# pre-built messages array (text + image_url parts) to the provider.
#
# If the gate module is not importable (exam-conductor not deployed),
# RuntimeError is raised — there is NO fallback to direct providers.
# The deployment must include exam-conductor for practice to function.
# ---------------------------------------------------------------------------

_gate_module = None          # lazily imported exam-conductor.llm_gate
_gate_import_attempted = False  # True once we have tried (success or failure)
_gate_unavailable = False    # True when import was attempted and failed


def _try_load_gate_module():
    """Attempt to import the LLM gate module once.  Returns the module or None."""
    global _gate_module, _gate_import_attempted, _gate_unavailable
    if _gate_module is not None:
        return _gate_module
    if _gate_import_attempted:
        return None
    _gate_import_attempted = True
    try:
        from api.v1._exampen_imports import load_exampen
        _gate_module = load_exampen("llm_gate")
        logger.info("LLM gate module loaded for practice bridge (SWM-011)")
        return _gate_module
    except ImportError:
        _gate_unavailable = True
        # CRITICAL: C4 requires all LLM calls through the gate.  If the gate
        # is not importable the deployment is non-compliant.
        logger.critical(
            "SWM-011 C4 VIOLATION: exam-conductor.llm_gate not importable — "
            "practice LLM calls (text + vision) will bypass the gate.  "
            "Deploy exam-conductor to restore compliance."
        )
        return None


async def _gate_text_call(
    db: DatabaseManager,
    current_user: Dict[str, Any],
    prompt: str,
    system_prompt: Optional[str] = None,
    max_tokens: int = 1000,
    temperature: float = 0.7,
    model_override: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """
    Route a text-only LLM call through the shared gate.

    Returns a dict shaped like ``AsyncOpenAIService.chat_completion_async``
    output (``success``, ``response``, etc.) or ``None`` **only** when the
    gate module was never importable (deployment issue).

    If the gate *is* loaded but the call fails at runtime, the exception
    propagates — there is no silent fallback to a direct provider.
    """
    gate_mod = _try_load_gate_module()
    if gate_mod is None:
        raise RuntimeError(
            "SWM-011 C4: LLM gate not available — exam-conductor not deployed. "
            "All LLM calls must go through the shared gate (C4)."
        )

    db_name = current_user.get("db_name")
    if not db_name:
        raise RuntimeError("SWM-011: No db_name in user context — cannot route through gate")

    tenant_db = await db.get_tenant_db(db_name)
    if tenant_db is None:
        raise RuntimeError(f"SWM-011: Could not obtain tenant DB '{db_name}' — cannot route through gate")

    gate = gate_mod.LLMGate(tenant_db)
    await gate.initialize()

    # Build the full prompt (system + user) since the gate takes a single prompt string
    full_prompt = prompt
    if system_prompt:
        full_prompt = f"{system_prompt}\n\n{prompt}"

    from config_async import OPENAI_MODEL
    model_id = model_override or OPENAI_MODEL

    gate_resp = await gate.call(
        model_id=model_id,
        prompt=full_prompt,
        caller_id="pcr_practice",
        max_output_tokens=max_tokens,
        temperature=temperature,
    )

    # Map GateResponse -> dict matching AsyncOpenAIService output shape
    return {
        "success": True,
        "response": gate_resp.content,
        "model": gate_resp.usage.model,
        "usage": {
            "prompt_tokens": gate_resp.usage.input_tokens,
            "completion_tokens": gate_resp.usage.output_tokens,
            "total_tokens": gate_resp.usage.total_tokens,
        },
    }

async def _gate_vision_call(
    db: DatabaseManager,
    current_user: Dict[str, Any],
    images: List[str],
    prompt: str,
    system_prompt: Optional[str] = None,
    max_tokens: int = 1000,
    temperature: float = 0.3,
    model_override: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """
    Route a vision/multimodal LLM call through the shared gate.

    Builds an OpenAI-style messages array with text + image_url parts and
    forwards it via ``gate.call(messages=...)``.

    Returns a dict shaped like ``AsyncOpenAIService.analyze_images_and_text_async``
    output (``success``, ``response``, etc.) or ``None`` **only** when the gate
    module was never importable (deployment issue).

    If the gate *is* loaded but the call fails at runtime, the exception
    propagates — there is no silent fallback to a direct provider.
    """
    gate_mod = _try_load_gate_module()
    if gate_mod is None:
        raise RuntimeError(
            "SWM-011 C4: LLM gate not available — exam-conductor not deployed. "
            "All LLM calls (including vision) must go through the shared gate (C4)."
        )

    db_name = current_user.get("db_name")
    if not db_name:
        raise RuntimeError("SWM-011: No db_name in user context — cannot route through gate")

    tenant_db = await db.get_tenant_db(db_name)
    if tenant_db is None:
        raise RuntimeError(f"SWM-011: Could not obtain tenant DB '{db_name}' — cannot route through gate")

    gate = gate_mod.LLMGate(tenant_db)
    await gate.initialize()

    # Build multimodal messages array (OpenAI format).
    # detail="high" forces the model to use its full ~2048x2048 internal representation
    # instead of "auto" (which downsamples to ~768x768 for sparse images). For handwritten
    # answers and chemistry/biology diagrams this is the difference between the model
    # reading individual strokes vs. seeing a blurred shape.
    content_parts: List[Dict[str, Any]] = [{"type": "text", "text": prompt}]
    for img in images or []:
        if not img:
            continue
        # Ensure data URI format
        if img.startswith("data:"):
            url = img
        else:
            url = f"data:image/png;base64,{img}"
        content_parts.append({
            "type": "image_url",
            "image_url": {"url": url, "detail": "high"},
        })

    messages: List[Dict[str, Any]] = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": content_parts})

    from config_async import OPENAI_MODEL
    model_id = model_override or OPENAI_MODEL

    gate_resp = await gate.call(
        model_id=model_id,
        prompt=prompt,
        caller_id="pcr_practice",
        messages=messages,
        max_output_tokens=max_tokens,
        temperature=temperature,
    )

    # Map GateResponse -> dict matching AsyncOpenAIService output shape
    return {
        "success": True,
        "response": gate_resp.content,
        "model": gate_resp.usage.model,
        "usage": {
            "prompt_tokens": gate_resp.usage.input_tokens,
            "completion_tokens": gate_resp.usage.output_tokens,
            "total_tokens": gate_resp.usage.total_tokens,
        },
    }


# Language detection utility for multilingual support
def detect_language(text: str) -> str:
    """
    Detect the primary language of the given text.
    
    Returns:
        'hindi' if the text contains significant Hindi (Devanagari) characters
        'english' otherwise (default)
    
    This is used to ensure LLM responses match the question/input language.
    """
    if not text:
        return 'english'
    
    # Count Devanagari characters (Hindi script range: U+0900 to U+097F)
    devanagari_count = sum(1 for char in text if '\u0900' <= char <= '\u097F')
    total_alpha_count = sum(1 for char in text if char.isalpha())
    
    if total_alpha_count == 0:
        return 'english'
    
    # If more than 20% of alphabetic characters are Devanagari, consider it Hindi
    hindi_ratio = devanagari_count / total_alpha_count
    
    if hindi_ratio > 0.2:
        return 'hindi'
    
    return 'english'


def robust_json_parse(raw_response: str) -> Optional[Dict[str, Any]]:
    """
    Robustly parse JSON from LLM response, handling:
    - Markdown code blocks
    - LaTeX content with escaped backslashes
    - Incomplete/truncated JSON
    - Extra text before/after JSON
    
    Returns parsed dict or None if parsing fails.
    """
    import json
    import re
    import ast
    
    if not raw_response:
        return None
    
    # Step 1: Remove markdown code blocks
    clean = raw_response.strip()
    
    # Remove ```json ... ``` blocks
    if "```json" in clean:
        match = re.search(r'```json\s*([\s\S]*?)\s*```', clean)
        if match:
            clean = match.group(1).strip()
        else:
            clean = re.sub(r'```json\s*', '', clean)
            clean = re.sub(r'\s*```', '', clean)
    elif "```" in clean:
        match = re.search(r'```\s*([\s\S]*?)\s*```', clean)
        if match:
            clean = match.group(1).strip()
        else:
            clean = re.sub(r'```\s*', '', clean)
    
    # Step 2: Try direct JSON parse
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        pass
    
    # Step 3: Find JSON object boundaries (handle nested braces)
    json_str = None
    
    # Find the first { and try to match balanced braces
    start_idx = clean.find('{')
    if start_idx != -1:
        brace_count = 0
        end_idx = start_idx
        in_string = False
        escape_next = False
        
        for i in range(start_idx, len(clean)):
            char = clean[i]
            
            if escape_next:
                escape_next = False
                continue
            
            if char == '\\':
                escape_next = True
                continue
            
            if char == '"' and not escape_next:
                in_string = not in_string
            
            if not in_string:
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        end_idx = i
                        break
        
        if brace_count == 0 and end_idx > start_idx:
            json_str = clean[start_idx:end_idx + 1]
    
    if json_str:
        # Try parsing the extracted JSON
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            pass
        
        # Step 4: Fix common JSON issues
        fixed = json_str
        
        # Fix unescaped control characters
        fixed = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', fixed)
        
        # Try again after fixing
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            pass
        
        # Step 5: Try ast.literal_eval (handles single quotes)
        try:
            result = ast.literal_eval(fixed)
            if isinstance(result, dict):
                return result
        except (ValueError, SyntaxError):
            pass
    
    # Step 6: Last resort - regex extraction for key fields
    # Try to extract individual fields if full parse fails
    try:
        extracted = {}

        # Helper: extract a JSON string value by key name.
        # Handles escaped quotes inside the value and stops at the
        # first unescaped closing quote.
        def _extract_string_field(key: str) -> Optional[str]:
            m = re.search(rf'"{key}"\s*:\s*"((?:[^"\\]|\\.)*)"', clean, re.DOTALL)
            if m:
                return m.group(1).replace('\\"', '"').replace('\\n', '\n')
            return None

        # Extract is_correct (boolean)
        is_correct_match = re.search(r'"is_correct"\s*:\s*(true|false)', clean, re.IGNORECASE)
        if is_correct_match:
            extracted["is_correct"] = is_correct_match.group(1).lower() == "true"

        # Extract score (number)
        score_match = re.search(r'"score"\s*:\s*([0-9.]+)', clean)
        if score_match:
            extracted["score"] = float(score_match.group(1))

        # Extract all string fields that the evaluation prompt requests
        _string_fields = [
            "extracted_answer", "solved_answer", "feedback", "reasoning",
            "work_shown", "what_went_wrong", "correct_solution",
            "solution", "final_answer",
        ]
        for field_name in _string_fields:
            val = _extract_string_field(field_name)
            if val:
                extracted[field_name] = val

        # If we got a meaningful subset, use the extracted data. This also helps
        # solution generation when a long model answer gets truncated before the
        # final JSON brace.
        if "is_correct" in extracted or "solution" in extracted:
            logger.info(f"📊 Partial JSON extraction succeeded: fields={list(extracted.keys())}")
            return extracted

    except Exception as e:
        logger.warning(f"Partial extraction failed: {e}")

    return None


def _normalize_latex_for_render(text: str) -> str:
    """Normalize LaTeX delimiters from any common dialect to the renderer's preferred form.

    The frontend's `VisualContentRenderer` accepts $...$, $$...$$, \\(...\\) and \\[...\\].
    LLMs are inconsistent — they sometimes mix dialects within a single response, or leave
    over-escaped backslashes from JSON-string handling. This helper normalises everything
    to $...$ / $$...$$ so rendering is uniform.

    Defensive only — when the prompt is followed correctly this is a no-op. When the LLM
    drifts, this catches it. Idempotent: applying twice produces the same result as once.
    """
    if not text:
        return text
    import re as _re
    out = text
    # Collapse over-escaped backslashes that survive a JSON round-trip:
    #   "\\\\frac" → "\\frac" → "\frac"
    # We only collapse the doubled variant; leaving single \\frac alone is correct LaTeX.
    out = out.replace("\\\\\\\\", "\\\\")
    # Convert display brackets \[ ... \] → $$ ... $$ (multi-line, non-greedy)
    out = _re.sub(r"\\\[([\s\S]*?)\\\]", r"$$\1$$", out)
    # Convert inline parens \( ... \) → $ ... $  (single-line preferred but allow multi-line)
    out = _re.sub(r"\\\(([\s\S]*?)\\\)", r"$\1$", out)
    try:
        from utils.latex_formatter import format_latex_in_text
        out = format_latex_in_text(out)
    except Exception as exc:
        logger.warning(f"Could not format raw LaTeX for rendering: {exc}")
    return out


def _clean_model_solution_text(text: str) -> str:
    """Convert JSON-ish cached model answers into display-ready text.

    The solution generator asks the LLM for JSON, but long case-study model
    answers can occasionally be returned or cached as the raw JSON envelope.
    This keeps that transport format out of the UI.
    """
    if not text:
        return ""

    import json as _json
    import re as _re

    raw = str(text).strip()
    if not raw:
        return ""
    raw = _re.sub(r"^```(?:json|markdown)?\s*", "", raw, flags=_re.IGNORECASE).strip()
    raw = _re.sub(r"\s*```$", "", raw).strip()

    parsed = robust_json_parse(raw)
    if isinstance(parsed, dict):
        candidate = (
            parsed.get("solution")
            or parsed.get("correct_solution")
            or parsed.get("model_answer")
            or parsed.get("answer")
        )
        if candidate:
            raw = str(candidate).strip()

    if raw.lstrip().startswith("{"):
        match = _re.search(r'"solution"\s*:\s*"((?:[^"\\]|\\.)*)"', raw, _re.DOTALL)
        if not match:
            match = _re.search(r'"correct_solution"\s*:\s*"((?:[^"\\]|\\.)*)"', raw, _re.DOTALL)
        if match:
            encoded = match.group(1)
            try:
                raw = _json.loads(f'"{encoded}"').strip()
            except Exception:
                raw = (
                    encoded
                    .replace('\\"', '"')
                    .replace("\\n", "\n")
                    .replace("\\t", "\t")
                    .replace("\\r", "")
                    .strip()
                )
        else:
            truncated = _re.search(r'"solution"\s*:\s*"', raw)
            if truncated:
                raw = raw[truncated.end():].strip()

    if len(raw) >= 2 and raw[0] == '"' and raw[-1] == '"':
        try:
            raw = _json.loads(raw).strip()
        except Exception:
            raw = raw[1:-1].strip()

    raw = (
        raw
        .replace("\\n", "\n")
        .replace("\\t", "\t")
        .replace("\\r", "")
        .replace('\\"', '"')
    )
    return _normalize_latex_for_render(raw.strip())


EVALUATION_MODE_STANDARD = "standard"
EVALUATION_MODE_CASE_STUDY = "case_study"
CASE_STUDY_SOLUTION_CACHE_VERSION = "case_study_compact_v2"


def _normalize_evaluation_mode(value: Any) -> Optional[str]:
    """Normalize an explicit per-question evaluation mode, if one is stored."""
    if value is None:
        return None
    raw = str(value).strip().lower().replace("-", "_").replace(" ", "_")
    if not raw:
        return None
    if raw in {"auto", "automatic", "default"}:
        return None
    if raw in {"case", "case_study", "business_case", "mba_case", "management_case"}:
        return EVALUATION_MODE_CASE_STUDY
    if raw in {"standard", "stem", "objective", "objective_stem", "subjective_stem"}:
        return EVALUATION_MODE_STANDARD
    return None


def _explicit_evaluation_mode(question_doc: Dict[str, Any]) -> Optional[str]:
    metadata = question_doc.get("metadata") or {}
    candidates = [
        question_doc.get("evaluation_mode"),
        question_doc.get("evaluationMode"),
        metadata.get("evaluation_mode") if isinstance(metadata, dict) else None,
        metadata.get("evaluationMode") if isinstance(metadata, dict) else None,
    ]
    for candidate in candidates:
        mode = _normalize_evaluation_mode(candidate)
        if mode:
            return mode
    return None


def _infer_evaluation_mode(question_doc: Dict[str, Any], *, is_mcq: Optional[bool] = None) -> str:
    """Infer grading style per question so mixed practice sets keep working."""
    explicit = _explicit_evaluation_mode(question_doc)
    if explicit:
        return explicit

    if is_mcq is None:
        stored_type = (question_doc.get("question_type") or "").lower().strip()
        is_mcq = bool(_options_text_from_question(question_doc)) and stored_type != "subjective"
    if is_mcq:
        return EVALUATION_MODE_STANDARD

    metadata = question_doc.get("metadata") or {}
    text_parts = [
        question_doc.get("text"),
        question_doc.get("question_text"),
        question_doc.get("subject"),
        question_doc.get("course_plan"),
        metadata.get("subject") if isinstance(metadata, dict) else "",
        metadata.get("course_plan") if isinstance(metadata, dict) else "",
    ]
    haystack = "\n".join(str(part or "") for part in text_parts).lower()

    business_terms = [
        "case study", "business case", "mba", "strategy", "strategic",
        "strategic analysis", "strategic considerations", "growth strategy",
        "growth lever", "growth levers", "core problem", "problem areas",
        "key metrics", "kpi", "90-day", "30/60/90", "go-to-market", "gtm",
        "unit economics", "cac", "ltv", "retention", "churn", "funnel",
        "segmentation", "pricing", "market share", "customer acquisition",
        "d2c", "brand", "sales have stagnated", "stagnated", "profitability",
        "business model", "operations", "supply chain", "founding team",
        "tam", "market sizing", "opportunity size", "pros/cons", "pros and cons",
        "swot", "porter", "five forces", "competition analysis", "competitive analysis",
        "risk", "risks", "business models", "recommendation", "should netflix",
        "live sports", "streaming rights", "sports rights",
    ]
    subject_terms = [
        "business", "management", "entrepreneurship", "marketing", "sales",
        "strategy", "strategic management",
    ]
    signal_count = sum(1 for term in business_terms if term in haystack)
    subject_signal = any(term in haystack for term in subject_terms)
    structured_prompt = (
        "identify:" in haystack
        or "recommend" in haystack
        or "should" in haystack
        or "strategy" in haystack
        or "pros" in haystack
        or "risks" in haystack
        or "\n-" in haystack
        or "\n*" in haystack
    )

    if signal_count >= 3:
        return EVALUATION_MODE_CASE_STUDY
    if signal_count >= 2 and structured_prompt:
        return EVALUATION_MODE_CASE_STUDY
    if subject_signal and signal_count >= 1 and structured_prompt:
        return EVALUATION_MODE_CASE_STUDY
    return EVALUATION_MODE_STANDARD


def _truncate_for_prompt(text: str, max_chars: int) -> str:
    """Trim large text blocks for prompt safety while preserving useful content."""
    if not text:
        return ""
    clean = str(text).strip()
    if len(clean) <= max_chars:
        return clean
    omitted = len(clean) - max_chars
    return f"{clean[:max_chars]}\n...[truncated {omitted} chars]"


router = APIRouter()

# Rate limiter
limiter = Limiter(key_func=get_remote_address)

# Pydantic models
class PracticeSession(BaseModel):
    id: Optional[str] = None
    student_id: str
    mode: str = Field(..., pattern="^(practice|exam|timed)$")
    subject: Optional[str] = None
    difficulty: Optional[str] = None
    questions_attempted: int = Field(default=0, ge=0)
    correct_answers: int = Field(default=0, ge=0)
    total_time_spent: int = Field(default=0, ge=0)  # in seconds
    started_at: datetime
    completed_at: Optional[datetime] = None
    is_completed: bool = False

class SessionQuestion(BaseModel):
    question_id: str
    answer: str
    is_correct: bool
    time_spent: int = Field(ge=0)  # in seconds
    answered_at: datetime

class SessionAnswer(BaseModel):
    question_id: str
    answer: str
    time_spent: int = Field(default=0, ge=0)

class StartSessionRequest(BaseModel):
    mode: str = Field(..., pattern="^(practice|exam|timed)$")
    subject: Optional[str] = None
    difficulty: Optional[str] = None
    time_limit: Optional[int] = Field(None, ge=1)  # in minutes
    document_id: Optional[str] = None  # Practice set document ID

class SessionResponse(BaseModel):
    id: str
    mode: str
    subject: Optional[str] = None
    difficulty: Optional[str] = None
    questions_attempted: int
    correct_answers: int
    accuracy_rate: float
    total_time_spent: int
    started_at: datetime
    completed_at: Optional[datetime] = None
    is_completed: bool

class SessionsListResponse(BaseModel):
    sessions: List[SessionResponse]
    total: int
    page: int
    limit: int

class PracticeStats(BaseModel):
    total_sessions: int
    total_time_spent: int
    average_accuracy: float
    sessions_by_mode: Dict[str, int]
    recent_activity: List[Dict[str, Any]]

# ----------------------
# Helper utilities (local)
# ----------------------

async def _extract_text_from_document(base64_data: str, doc_type: str, filename: str) -> str:
    """Extract text content from uploaded PDF or DOCX document.
    
    Args:
        base64_data: Base64 encoded document data (may include data URL prefix)
        doc_type: Type of document ('pdf' or 'docx')
        filename: Original filename for logging
        
    Returns:
        Extracted text content from the document
    """
    import base64
    import io
    import tempfile
    import os
    
    try:
        # Remove data URL prefix if present
        if ',' in base64_data:
            base64_data = base64_data.split(',')[-1]
        
        # Decode base64 to bytes
        doc_bytes = base64.b64decode(base64_data)
        
        if doc_type == 'pdf':
            try:
                from pypdf import PdfReader
                pdf_file = io.BytesIO(doc_bytes)
                pdf_reader = PdfReader(pdf_file)
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                
                return "\n\n".join(text_content)
            except ImportError:
                logger.error("pypdf not installed. Cannot extract PDF text. Install with: pip install pypdf")
                return ""
        
        elif doc_type == 'docx':
            try:
                from docx import Document
                docx_file = io.BytesIO(doc_bytes)
                doc = Document(docx_file)
                
                text_content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                return "\n".join(text_content)
            except ImportError:
                logger.error("python-docx not installed. Cannot extract DOCX text.")
                return ""
        
        else:
            logger.warning(f"Unsupported document type: {doc_type}")
            return ""
            
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        return ""


async def _load_question_doc(db: DatabaseManager, qid: str, is_b2c: bool = False) -> Dict[str, Any]:
    """Fetch question from MongoDB.
    
    For B2C users, uses B2C database instead of main database.
    """
    if is_b2c:
        return await db.b2c_find_one("questions", {"id": qid}) or {}
    return await db.mongo_find_one("questions", {"id": qid}) or {}


def _question_content_hash(question_doc: Dict[str, Any], evaluation_mode: Optional[str] = None) -> str:
    """Stable hash of the question's *content* for solution-cache invalidation.

    Hashes whatever a tutor would consider "the question": text, options,
    admin-provided correct answer, and the identity (not full bytes) of each
    attached figure. If the admin edits any of these, the hash changes and a
    fresh solution gets generated next time a student attempts this question.
    """
    import hashlib
    import json as _json

    mode = (
        evaluation_mode
        or question_doc.get("evaluation_mode")
        or question_doc.get("evaluationMode")
        or ""
    )
    parts: List[str] = [
        str(question_doc.get("text") or "").strip(),
        str(question_doc.get("correctAnswer") or question_doc.get("correct_answer") or "").strip(),
        str(mode).strip(),
    ]
    if _normalize_evaluation_mode(mode) == EVALUATION_MODE_CASE_STUDY:
        parts.append(CASE_STUDY_SOLUTION_CACHE_VERSION)

    opts = question_doc.get("options") or []
    if opts:
        parts.append(_json.dumps([str(o) for o in opts], ensure_ascii=False, sort_keys=True))
    enh = question_doc.get("enhancedOptions") or []
    if enh:
        compact = [
            {"type": (o or {}).get("type"), "content": (o or {}).get("content"), "label": (o or {}).get("label")}
            for o in enh
        ]
        parts.append(_json.dumps(compact, ensure_ascii=False, sort_keys=True))

    # Identify figures by id + size only (not full base64 payload — too expensive).
    fig_summary: List[Dict[str, Any]] = []
    for fig in question_doc.get("questionFigures") or []:
        fig_summary.append({
            "id": (fig or {}).get("id"),
            "url": (fig or {}).get("url"),
            "path": (fig or {}).get("path"),
            "len": len((fig or {}).get("base64Data") or "") or None,
        })
    for img in question_doc.get("images") or []:
        if (img or {}).get("type") == "diagram":
            fig_summary.append({
                "id": (img or {}).get("id"),
                "url": (img or {}).get("url"),
                "path": (img or {}).get("path"),
                "len": len((img or {}).get("base64Data") or "") or None,
            })
    if fig_summary:
        parts.append(_json.dumps(fig_summary, ensure_ascii=False, sort_keys=True))

    canonical = "||".join(parts)
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()[:32]


async def _get_or_generate_solution(
    question_doc: Dict[str, Any],
    db: DatabaseManager,
    current_user: Dict[str, Any],
    is_b2c: bool,
    evaluation_mode: Optional[str] = None,
) -> Dict[str, str]:
    """Return the canonical step-by-step solution for a question.

    Strategy:
      1. If `solution_cache` exists on the question doc AND its `contentHash`
         matches the question's current content → use the cache (1 DB read,
         0 LLM calls).
      2. Else generate the solution via a single LLM call that sees ONLY the
         question (text + options + figures + admin's correct_answer).
         Student canvas pages are deliberately excluded so a student's wrong
         attempt cannot pollute the shared solution.
      3. Save the result back to the `questions` collection so every future
         student gets the same solution for free.

    Returns a dict with keys: solution (str), final_answer (str), source (str).
    """
    qid = question_doc.get("id") or ""
    mode = _normalize_evaluation_mode(evaluation_mode) or _infer_evaluation_mode(question_doc)
    current_hash = _question_content_hash(question_doc, mode)
    cache = question_doc.get("solution_cache") or {}

    if (
        isinstance(cache, dict)
        and cache.get("contentHash") == current_hash
        and (cache.get("correctSolution") or "").strip()
    ):
        cached_solution_raw = str(cache.get("correctSolution") or "").strip()
        cached_solution_text = _clean_model_solution_text(cached_solution_raw)
        if cached_solution_text != cached_solution_raw:
            try:
                update = {"$set": {"solution_cache.correctSolution": cached_solution_text}}
                if is_b2c:
                    await db.b2c_update_one("questions", {"id": qid}, update)
                else:
                    await db.mongo_update_one("questions", {"id": qid}, update)
                logger.info(f"🧹 Cleaned cached solution display text for Q:{qid}")
            except Exception as clean_err:
                logger.warning(f"Could not clean persisted solution cache for Q:{qid}: {clean_err}")
        logger.info(f"📚 Solution cache HIT for Q:{qid} (source={cache.get('source')})")
        return {
            # Normalise on read too — old cached entries may pre-date the normaliser.
            "solution": cached_solution_text,
            "final_answer": _normalize_latex_for_render(str(cache.get("finalAnswer") or "").strip()),
            "source": str(cache.get("source") or "cache"),
        }

    logger.info(
        f"📚 Solution cache MISS for Q:{qid} "
        f"(have={bool(cache.get('correctSolution'))}, hash_match="
        f"{cache.get('contentHash') == current_hash if cache else False}). Generating..."
    )

    # Build solution-only prompt — no student work allowed in here.
    question_text = str(question_doc.get("text") or "")
    options_text = _options_text_from_question(question_doc)
    ca_primary = question_doc.get("correctAnswer")
    ca_alt = question_doc.get("correct_answer")
    correct_answer_raw = str(
        ca_primary if ca_primary is not None else (ca_alt if ca_alt is not None else "")
    ).strip()

    # Resolve the admin's stored answer: if it's an option letter, look up what that
    # letter actually represents. We pass the *value* to the LLM, not just the letter,
    # so the LLM can never produce the "I derived 2mv₀² but I'll label it option D"
    # contradiction we see when only the letter is shown.
    resolved = _resolve_correct_answer(correct_answer_raw, question_doc) if correct_answer_raw else {
        "raw": "", "resolved_value": "", "display": "", "is_option_letter": False,
    }
    admin_letter = resolved["raw"]
    admin_value = (resolved["resolved_value"] or "").strip()
    admin_is_letter = bool(resolved["is_option_letter"])

    if mode == EVALUATION_MODE_CASE_STUDY:
        prompt_parts: List[str] = [
            "Generate a compact MBA-style reference framework for the following case-study "
            "question. This will be shown to every student who attempts the question, so it "
            "must be self-contained and must not reference any specific student attempt. "
            "Do not write an exhaustive essay. Keep the answer practical, concise, and tied "
            "to the case facts.",
            "",
            "QUESTION:",
            question_text,
        ]
        if options_text:
            prompt_parts.append("\nOPTIONS / CASE MATERIAL:")
            prompt_parts.append(options_text)

        if correct_answer_raw:
            prompt_parts.append("\nTEACHER'S REFERENCE ANSWER OR RUBRIC:")
            prompt_parts.append(admin_value or admin_letter)
            prompt_parts.append(
                "Use this as grading guidance, but do not require the student to match "
                "the wording exactly. Case-study answers can be correct through sound "
                "business reasoning, prioritization, and execution logic."
            )
        else:
            prompt_parts.append(
                "\nNo teacher rubric was provided. Create a strong reference answer using "
                "standard MBA case-analysis expectations."
            )

        prompt_parts.append('''
Return strict JSON (no markdown fences, no commentary):
{
  "solution": "a compact MBA-style reference framework under 500 words. Use exactly these sections: Executive Diagnosis, Core Problem Areas, Growth Levers, Key Metrics, 90-Day Plan, Risks/Assumptions. Use 1-3 concise bullets per section. Do not invent precise numbers that are not in the case; use assumptions where needed.",
  "final_answer": "a short summary of what a strong answer should include",
  "teacher_answer_disagreement": false
}
For case studies, set `teacher_answer_disagreement` to false unless the teacher's reference answer is internally impossible or contradicts the case facts.'''.strip())

        system_prompt = (
            "You are an MBA case evaluator writing a reusable model answer and rubric. "
            "Favor structured business reasoning over one exact answer. Keep the reference "
            "answer compact: no more than 500 words, no exhaustive lists, no unsupported "
            "precise numbers, no speculation beyond clear assumptions. Include practical "
            "diagnosis, prioritized actions, metrics, timeline, assumptions, and risks. "
            "Output ONLY valid JSON, no markdown code fences, no commentary."
        )
    else:
        prompt_parts = [
            "Generate a clear, complete model solution to the following question. "
            "This solution will be shown to every student who attempts the question, so it must be "
            "self-contained and must not reference any specific student attempt.",
            "",
            "QUESTION:",
            question_text,
        ]
        if options_text:
            prompt_parts.append("\nOPTIONS:")
            prompt_parts.append(options_text)

        if correct_answer_raw:
            prompt_parts.append("\nTEACHER'S STATED ANSWER:")
            if admin_is_letter and admin_value and admin_value != admin_letter:
                prompt_parts.append(
                    f"The teacher marked option {admin_letter}, which represents: {admin_value}"
                )
                prompt_parts.append(
                    f"\nSolve the question yourself, step by step. Your derivation must arrive at the "
                    f"VALUE {admin_value} (not just label your answer with the letter '{admin_letter}'). "
                    f"If your honest derivation lands on a different value, set "
                    f"`teacher_answer_disagreement` to true in the JSON and report what you actually "
                    f"derived - do NOT force-fit the math to match {admin_value}."
                )
            else:
                prompt_parts.append(admin_value or admin_letter)
                prompt_parts.append(
                    "\nSolve the question yourself, step by step. Your derivation must arrive at this "
                    "answer. If your honest derivation produces a different answer, set "
                    "`teacher_answer_disagreement` to true and report what you actually derived."
                )
        else:
            prompt_parts.append(
                "\nNo answer was provided by the teacher - solve the question yourself and write the "
                "step-by-step derivation."
            )

        prompt_parts.append('''
Return strict JSON (no markdown fences, no commentary):
{
  "solution": "the step-by-step solution, written for a student. Use $...$ for inline math or chemistry notation (e.g. $x_0$, $\\frac{a}{b}$, $\\ce{H2O}$) and $$...$$ for display math. Do NOT use \\(...\\) or \\[...\\]. Formulas and identifiers must be wrapped in $...$.",
  "final_answer": "the short final answer this derivation produces (e.g. \\"$\\frac{7}{12}$\\", \\"$\\ce{H2O}$\\", \\"D\\", \\"Delhi\\"). Wrap any math or chemistry notation in $...$.",
  "teacher_answer_disagreement": false
}
Set `teacher_answer_disagreement` to true ONLY if your honest derivation lands on a value different from the teacher's stated answer. Do not flip the field for minor formatting differences (e.g. \\"7 days\\" vs \\"7\\" when the question is about days). When unsure, default to false.'''.strip())

        system_prompt = (
            "You are an expert tutor writing a model solution. The solution must be academically "
            "correct, concise, and pedagogical. It will be cached and shown to every student who "
            "attempts this question, so do NOT mention any specific student or attempt. "
            "Use $...$ for inline math or chemistry notation and $$...$$ for display math; never "
            "\\(...\\) or \\[...\\]. Wrap variables, fractions, formulas, and chemistry expressions "
            "such as $\\ce{H2O}$ in delimiters so the renderer typesets them correctly. "
            "Output ONLY valid JSON, no markdown code fences, no commentary."
        )

    full_prompt = "\n".join(prompt_parts)
    question_images = await _figure_images_base64(question_doc, db, is_b2c)
    option_images_data = await _option_images_base64(question_doc, db, is_b2c)
    option_images = [oi["image"] for oi in option_images_data]
    all_question_images = question_images + option_images

    try:
        generation_max_tokens = 1400 if mode == EVALUATION_MODE_CASE_STUDY else 1500
        if all_question_images:
            response = await _gate_vision_call(
                db, current_user, all_question_images, full_prompt,
                system_prompt=system_prompt,
                max_tokens=generation_max_tokens,
                temperature=0.2,
            )
        else:
            response = await _gate_text_call(
                db, current_user, full_prompt,
                system_prompt=system_prompt,
                max_tokens=generation_max_tokens,
                temperature=0.2,
            )
        raw = (response.get("response") or "").strip()
        parsed = robust_json_parse(raw) or {}
        solution_text = _clean_model_solution_text(str(parsed.get("solution") or raw).strip())
        final_answer = _normalize_latex_for_render(str(parsed.get("final_answer") or "").strip())
        llm_flagged_disagreement = bool(parsed.get("teacher_answer_disagreement"))

        if not solution_text:
            # Fallback: use raw response so we at least have *something* useful.
            logger.warning(f"⚠️ Solution generation produced no JSON for Q:{qid}; using raw text.")
            solution_text = _clean_model_solution_text(raw[:8000])
    except Exception as gen_err:
        logger.error(f"❌ Solution generation failed for Q:{qid}: {gen_err}", exc_info=True)
        return {"solution": "", "final_answer": correct_answer_raw or "", "source": "error"}

    # Detect admin/LLM disagreement, even when the LLM forgot to flag it. Use the existing
    # _answers_are_equivalent helper if present, else a simple normalised compare.
    admin_llm_disagree = False
    if mode != EVALUATION_MODE_CASE_STUDY and admin_value and final_answer:
        s_norm = "".join(final_answer.split()).lower()
        a_norm = "".join(admin_value.split()).lower()
        # Strip common LaTeX wrapping and punctuation so "$2mv_0^2$" matches "2mv_0^2".
        for ch in ("$", "\\(", "\\)", "\\[", "\\]", "(", ")", " ", ",", "."):
            s_norm = s_norm.replace(ch, "")
            a_norm = a_norm.replace(ch, "")
        if s_norm and a_norm and s_norm != a_norm:
            admin_llm_disagree = True
    if mode != EVALUATION_MODE_CASE_STUDY and llm_flagged_disagreement:
        admin_llm_disagree = True

    if admin_llm_disagree:
        logger.warning(
            f"⚠️ ADMIN/LLM ANSWER MISMATCH on Q:{qid} — "
            f"admin stated '{admin_letter}' (value='{admin_value}') but solution derives "
            f"'{final_answer}'. Solution stored is what the LLM actually derived, NOT a forced "
            f"fit to the admin answer. Question should be reviewed."
        )

    cache_doc = {
        "correctSolution": solution_text,
        "finalAnswer": final_answer or admin_value or correct_answer_raw,
        "adminStatedAnswer": admin_letter,
        "adminStatedValue": admin_value,
        "adminLlmDisagree": admin_llm_disagree,
        "source": "admin" if (correct_answer_raw and not admin_llm_disagree) else "llm_generated",
        "evaluationMode": mode,
        "solutionCacheVersion": CASE_STUDY_SOLUTION_CACHE_VERSION if mode == EVALUATION_MODE_CASE_STUDY else "standard",
        "contentHash": current_hash,
        "model": str(response.get("model") or ""),
        "updatedAt": datetime.utcnow(),
    }

    try:
        if is_b2c:
            await db.b2c_update_one(
                "questions", {"id": qid}, {"$set": {"solution_cache": cache_doc}}
            )
        else:
            await db.mongo_update_one(
                "questions", {"id": qid}, {"$set": {"solution_cache": cache_doc}}
            )
        logger.info(f"💾 Cached solution for Q:{qid} (source={cache_doc['source']}, len={len(solution_text)})")
    except Exception as save_err:
        logger.error(f"❌ Failed to persist solution_cache for Q:{qid}: {save_err}")

    return {
        "solution": solution_text,
        "final_answer": cache_doc["finalAnswer"],
        "source": cache_doc["source"],
    }


def _options_text_from_question(q: Dict[str, Any]) -> str:
    opts = q.get("options", []) or []
    if opts:
        return "\n".join([f"{chr(65+i)}. {opt}" for i, opt in enumerate(opts)])
    enh = q.get("enhancedOptions") or []
    if enh:
        parts = []
        for i, opt in enumerate(enh):
            label = chr(65 + i)
            if isinstance(opt, dict):
                if opt.get("type") == "text" and opt.get("content"):
                    parts.append(f"{label}. {opt.get('content')}")
                elif opt.get("type") == "image":
                    desc = opt.get("description") or "image option"
                    parts.append(f"{label}. [IMAGE] {desc}")
            else:
                parts.append(f"{label}. {str(opt)}")
        return "\n".join(parts)
    return ""


def _resolve_correct_answer(correct_answer: str, question_doc: dict) -> dict:
    """
    Resolve the correct answer into multiple representations for robust evaluation.

    Returns dict with:
        - raw: original stored value (e.g. "A")
        - resolved_value: actual content of the option (e.g. "7" for option A)
        - display: human-readable for summary (e.g. "A (7)")
        - is_option_letter: True if the answer is A/B/C/D
    """
    result = {
        "raw": correct_answer,
        "resolved_value": correct_answer,
        "display": correct_answer,
        "is_option_letter": False,
    }

    if not correct_answer:
        return result

    ca_upper = correct_answer.strip().upper()

    # Check if it's an option letter (A-J covers most MCQ ranges)
    if len(ca_upper) == 1 and ca_upper in "ABCDEFGHIJ":
        option_index = ord(ca_upper) - ord('A')

        # Try to resolve from enhancedOptions first, then plain options
        enhanced = question_doc.get("enhancedOptions") or []
        opts = question_doc.get("options", []) or []

        resolved_content = None
        if enhanced and option_index < len(enhanced):
            opt = enhanced[option_index]
            if isinstance(opt, dict):
                content = opt.get("content", "")
                if content and content.strip():
                    resolved_content = content.strip()
        elif opts and option_index < len(opts):
            if opts[option_index] and str(opts[option_index]).strip():
                resolved_content = str(opts[option_index]).strip()

        if resolved_content:
            result["is_option_letter"] = True
            result["resolved_value"] = resolved_content
            result["display"] = f"{ca_upper} ({resolved_content})"
        else:
            # It looks like a letter but we couldn't resolve it — still mark as option letter
            result["is_option_letter"] = True

    return result


async def _figure_images_base64(q: Dict[str, Any], db: DatabaseManager = None, is_b2c: bool = False) -> List[str]:
    """Extract base64 image data for question figures.
    
    ENHANCED: Now loads images from disk/database if base64Data is not embedded.
    This ensures question diagram images are always available for LLM evaluation.
    
    Args:
        q: Question document
        db: Database manager for loading images from disk (optional)
        is_b2c: Whether this is a B2C user (uses B2C database)
        
    Returns:
        List of base64 data URLs for question figures
    """
    import os
    import base64 as base64_module
    
    imgs: List[str] = []
    
    # Helper function to load image from database/disk
    async def load_image_by_id(img_id: str) -> Optional[str]:
        if not img_id or not db:
            return None
        try:
            if is_b2c:
                img_doc = await db.b2c_find_one("images", {"_id": img_id})
            else:
                img_doc = await db.mongo_find_one("images", {"_id": img_id})
            
            if img_doc:
                if img_doc.get("base64Data"):
                    b64 = img_doc["base64Data"]
                    logger.info(f"Loaded base64Data from database for image {img_id}")
                    return b64
                elif img_doc.get("file_path"):
                    file_path = img_doc["file_path"]
                    if os.path.exists(file_path):
                        async with aiofiles.open(file_path, "rb") as f:
                            image_bytes = await f.read()
                        base64_encoded = base64_module.b64encode(image_bytes).decode('utf-8')
                        content_type = img_doc.get("content_type", "image/jpeg")
                        if not content_type.startswith("image/"):
                            content_type = "image/jpeg"
                        b64 = f"data:{content_type};base64,{base64_encoded}"
                        logger.info(f"Loaded image from disk for {img_id}: {len(b64)} bytes")
                        return b64
        except Exception as e:
            logger.error(f"Failed to load image {img_id}: {e}")
        return None
    
    # Helper to normalize base64 format
    def normalize_b64(b64: str) -> str:
        if b64 and not b64.startswith("data:image"):
            return f"data:image/png;base64,{b64}"
        return b64
    
    # 1. Extract from question_figures array
    for fig_ref in (q.get("question_figures", []) or []):
        try:
            b64 = None
            fig_id = None
            
            if isinstance(fig_ref, dict):
                b64 = fig_ref.get("base64Data")
                fig_id = fig_ref.get("id")
            elif isinstance(fig_ref, str):
                fig_id = fig_ref
            
            if not b64 and fig_id:
                b64 = await load_image_by_id(fig_id)
            
            if b64:
                imgs.append(normalize_b64(b64))
                
        except Exception as e:
            logger.warning(f"Error processing figure: {e}")
    
    # 2. Extract from images array (diagrams and other embedded images)
    for img_ref in (q.get("images", []) or []):
        try:
            b64 = None
            img_id = None
            
            if isinstance(img_ref, dict):
                # Include ALL image types, not just diagrams
                b64 = img_ref.get("base64Data")
                img_id = img_ref.get("id")
            elif isinstance(img_ref, str):
                img_id = img_ref
            
            if not b64 and img_id:
                b64 = await load_image_by_id(img_id)
            
            if b64:
                imgs.append(normalize_b64(b64))
        except Exception:
            pass
    
    # 3. Check for inline figure in the question text (base64 embedded)
    if q.get("figure") and isinstance(q.get("figure"), str):
        fig = q["figure"]
        if fig.startswith("data:image") or len(fig) > 100:  # Likely base64
            imgs.append(normalize_b64(fig))
    
    # 4. Check for questionImage field
    if q.get("questionImage"):
        qimg = q["questionImage"]
        if isinstance(qimg, str):
            if qimg.startswith("data:image") or len(qimg) > 100:
                imgs.append(normalize_b64(qimg))
            else:
                # It might be an ID
                loaded = await load_image_by_id(qimg)
                if loaded:
                    imgs.append(normalize_b64(loaded))
        elif isinstance(qimg, dict):
            b64 = qimg.get("base64Data")
            if not b64:
                b64 = await load_image_by_id(qimg.get("id"))
            if b64:
                imgs.append(normalize_b64(b64))
    
    logger.info(f"📷 Extracted {len(imgs)} question figure/diagram images for LLM evaluation")
    return imgs


async def _option_images_base64(q: Dict[str, Any], db: DatabaseManager = None, is_b2c: bool = False) -> List[Dict[str, str]]:
    """Extract images from MCQ options.
    
    Returns list of dicts with 'option' (A, B, C, D) and 'image' (base64 data URL)
    """
    import os
    import base64 as base64_module
    
    option_images: List[Dict[str, str]] = []
    
    # Helper to load image by ID
    async def load_image_by_id(img_id: str) -> Optional[str]:
        if not img_id or not db:
            return None
        try:
            if is_b2c:
                img_doc = await db.b2c_find_one("images", {"_id": img_id})
            else:
                img_doc = await db.mongo_find_one("images", {"_id": img_id})
            
            if img_doc:
                if img_doc.get("base64Data"):
                    return img_doc["base64Data"]
                elif img_doc.get("file_path"):
                    file_path = img_doc["file_path"]
                    if os.path.exists(file_path):
                        async with aiofiles.open(file_path, "rb") as f:
                            image_bytes = await f.read()
                        base64_encoded = base64_module.b64encode(image_bytes).decode('utf-8')
                        content_type = img_doc.get("content_type", "image/jpeg")
                        return f"data:{content_type};base64,{base64_encoded}"
        except Exception as e:
            logger.error(f"Failed to load option image {img_id}: {e}")
        return None
    
    def normalize_b64(b64: str) -> str:
        if b64 and not b64.startswith("data:image"):
            return f"data:image/png;base64,{b64}"
        return b64
    
    # Check enhancedOptions for image options
    enh = q.get("enhancedOptions") or []
    for i, opt in enumerate(enh):
        label = chr(65 + i)  # A, B, C, D...
        if isinstance(opt, dict) and opt.get("type") == "image":
            b64 = opt.get("base64Data")
            if not b64:
                b64 = await load_image_by_id(opt.get("id"))
            if b64:
                option_images.append({
                    "option": label,
                    "image": normalize_b64(b64)
                })
    
    # Also check regular options array for image objects
    opts = q.get("options") or []
    for i, opt in enumerate(opts):
        label = chr(65 + i)
        if isinstance(opt, dict):
            if opt.get("type") == "image" or opt.get("image"):
                b64 = opt.get("base64Data") or opt.get("image")
                if not b64:
                    b64 = await load_image_by_id(opt.get("id"))
                if b64:
                    option_images.append({
                        "option": label,
                        "image": normalize_b64(b64)
                    })
    
    if option_images:
        logger.info(f"📷 Extracted {len(option_images)} option images: {[o['option'] for o in option_images]}")
    
    return option_images

def _normalize_choice_text(s: str) -> str:
    import re as _re
    t = (s or '').upper().strip()
    # Support any single letter A-Z for MCQ options (not just A-D)
    m = _re.search(r"\b([A-Z])\b", t)
    return m.group(1) if m else t

def _normalize_numeric_text(s: str) -> str:
    t = (s or '').strip().replace(' ', '').replace(',', '.')
    if ':' in t and t.count(':') == 1 and all(part.isdigit() for part in t.split(':')):
        t = t.replace(':', '.')
    return t

def require_student_or_admin(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Dependency to require student, admin, or B2C user access"""
    allowed_types = ["student", "admin", "b2c_user", "b2c_admin"]
    if current_user.get("user_type") not in allowed_types:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Student, admin, or B2C user access required"
        )
    return current_user

@router.post("/next")
@limiter.limit("60/minute")
async def get_next_practice_question(
    request: Request,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    """Return a random next question from MongoDB.

    Behavior:
    - Prefer questions tagged as Practice Sets in metadata when available
    - Fall back gracefully to all questions if tag is missing
    - Build figure/image payloads with base64 when available; else serve via images API
    """
    try:
        import random
        from pydantic import BaseModel
        
        class NextQuestionRequest(BaseModel):
            subject: Optional[str] = None
            difficulty: Optional[str] = None
            excludeIds: Optional[List[str]] = None

        # Safely parse request body (optional)
        subject: Optional[str] = None
        difficulty: Optional[str] = None
        exclude_ids: List[str] = []
        try:
            if request.headers.get("content-type", "").startswith("application/json"):
                body = await request.json()
                if isinstance(body, dict):
                    req_data = NextQuestionRequest(**body)
                    subject = req_data.subject
                    difficulty = req_data.difficulty
                    if req_data.excludeIds:
                        exclude_ids = list(req_data.excludeIds)
        except Exception as _e:
            # Fall back to defaults if parsing fails; do not reject
            exclude_ids = []
        
        # Get questions from MongoDB (Practice Sets)
        fetched_ids = []
        metadatas = []

        admin_id = current_user.get("admin_id")
        admin_id_str = str(admin_id).strip() if admin_id is not None else ""

        def _apply_admin_scope(query_filter: Dict[str, Any]) -> None:
            """Attach admin isolation filter when available."""
            if not admin_id_str:
                return
            try:
                query_filter["admin_id"] = ObjectId(admin_id_str)
            except Exception:
                query_filter["admin_id"] = admin_id_str
        
        # Build filter for Practice Sets
        mongo_filter: Dict[str, Any] = {"document_type": "Practice Sets"}
        _apply_admin_scope(mongo_filter)
        if subject:
            mongo_filter["subject"] = subject
        if difficulty:
            mongo_filter["difficulty"] = difficulty

        mongo_questions = await db.mongo_find("questions", mongo_filter, limit=1000)
        fetched_ids = [q.get("id") for q in mongo_questions if q.get("id")]
        
        for q in mongo_questions:
            metadata = {
                "fullData": json.dumps(q, default=str),
                "subject": q.get("subject", ""),
                "difficulty": q.get("difficulty", "medium"),
                "document_type": q.get("document_type", "Practice Sets")
            }
            metadatas.append(metadata)
            
        logger.info(f"Fetched {len(fetched_ids)} Practice Sets questions from MongoDB (subject={subject}, difficulty={difficulty})")

        # Additional MongoDB fallback if no results
        if not fetched_ids:
            mongo_filter = {"metadata.document_type": "Practice Sets"}
            _apply_admin_scope(mongo_filter)
            if subject:
                mongo_filter["subject"] = subject
            if difficulty:
                mongo_filter["difficulty"] = difficulty

            mongo_questions = await db.mongo_find("questions", mongo_filter, projection={"id": 1})
            fetched_ids = [q.get("id") for q in mongo_questions if q.get("id")]
            logger.info(f"MongoDB fallback fetched {len(fetched_ids)} question ids")

        if not fetched_ids:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No practice questions found. Please upload Practice Sets documents and process them."
            )

        # Refine the pool via fullData if available
        if metadatas and fetched_ids:
            refined: List[str] = []
            for qid, md in zip(fetched_ids, metadatas):
                full_json = md.get('fullData')
                if not full_json:
                    refined.append(qid)
                    continue
                try:
                    import json as _json
                    fd = _json.loads(full_json)
                    doc_type = (fd.get('metadata', {}) or {}).get('document_type')
                    if doc_type and doc_type != 'Practice Sets':
                        continue
                    if subject and fd.get('subject') != subject:
                        continue
                    if difficulty and fd.get('difficulty') != difficulty:
                        continue
                    refined.append(qid)
                except Exception:
                    refined.append(qid)
            fetched_ids = refined

        # Filter out excluded IDs
        available_ids = [qid for qid in fetched_ids if qid not in exclude_ids]
        
        if not available_ids:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No new questions available. All questions have been attempted."
            )
        
        # Select random question from available
        question_id = random.choice(available_ids)
        
        # Get question from MongoDB
        question_doc = await db.mongo_find_one("questions", {"id": question_id}) or {}
        
        if not question_doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Question data not found"
            )
        
        # Build image URLs for frontend - images are served from local disk via /api/v1/images/{id}
        images_with_urls = []
        for img_ref in question_doc.get("images", []) or []:
            # Support both string ID and dict refs having id
            img_id = img_ref.get("id") if isinstance(img_ref, dict) else img_ref
            if not img_id:
                continue
            # Check if image exists in MongoDB using _id field (MongoDB primary key)
            img_doc = await db.mongo_find_one("images", {"_id": img_id})
            if img_doc:
                images_with_urls.append({
                    "id": img_id,
                    "url": f"/api/v1/images/{img_id}",  # Serve from local disk
                    "contentType": img_doc.get("content_type", "image/jpeg"),
                    "filename": img_doc.get("original_filename", str(img_id))
                })
            else:
                logger.warning(f"Image {img_id} referenced in question {question_id} but not found in database")
        
        # Also include QUESTION FIGURES (diagrams)
        figures_with_urls = []
        for fig_ref in question_doc.get("question_figures", []):
            try:
                fig_id = fig_ref.get("id") if isinstance(fig_ref, dict) else fig_ref
                base64_data = None

                # First check if base64Data is embedded in the figure reference
                if isinstance(fig_ref, dict) and fig_ref.get("base64Data"):
                    base64_data = fig_ref["base64Data"]
                    logger.info(f"Using embedded base64Data for figure {fig_id}")
                else:
                    # Try to get base64Data from images collection
                    img_doc = await db.mongo_find_one("images", {"_id": fig_id})
                    if img_doc:
                        # Check if base64Data is stored in the document
                        if img_doc.get("base64Data"):
                            base64_data = img_doc["base64Data"]
                            logger.info(f"Using stored base64Data for figure {fig_id}")
                        # If not, read from file_path and convert to base64
                        elif img_doc.get("file_path"):
                            import os
                            import base64
                            file_path = img_doc["file_path"]
                            if os.path.exists(file_path):
                                try:
                                    async with aiofiles.open(file_path, "rb") as f:
                                        image_bytes = await f.read()
                                    base64_encoded = base64.b64encode(image_bytes).decode('utf-8')
                                    # Determine content type from file extension or stored content_type
                                    content_type = img_doc.get("content_type", "image/jpeg")
                                    if not content_type.startswith("image/"):
                                        # Default to jpeg if content type is not an image
                                        content_type = "image/jpeg"
                                    base64_data = f"data:{content_type};base64,{base64_encoded}"
                                    logger.info(f"✅ Loaded and converted image {fig_id} from file: {len(base64_data)} bytes")
                                except Exception as file_err:
                                    logger.error(f"❌ Failed to read image file {file_path}: {file_err}")
                            else:
                                logger.warning(f"⚠️ Image file not found: {file_path}")
                        else:
                            logger.warning(f"⚠️ No base64Data or file_path for image {fig_id}")
                    else:
                        logger.warning(f"⚠️ Image document not found: {fig_id}")

                figures_with_urls.append({
                    "id": fig_id,
                    "url": f"/api/v1/images/{fig_id}",
                    "contentType": "image/jpeg",
                    "filename": (fig_ref.get("filename") if isinstance(fig_ref, dict) else str(fig_id)),
                    "base64Data": base64_data,
                    "description": (fig_ref.get("description", "") if isinstance(fig_ref, dict) else ""),
                    "type": "diagram"
                })
            except Exception as _e:
                logger.error(f"❌ Practice figures processing error: {_e}", exc_info=True)

        merged_images = images_with_urls + figures_with_urls

        # Format LaTeX in question text and options
        from utils.latex_formatter import format_question_latex

        evaluation_mode = _infer_evaluation_mode(question_doc)
        question = {
            "id": question_id,
            "text": question_doc.get("text", ""),
            "subject": question_doc.get("subject", ""),
            "difficulty": question_doc.get("difficulty", "medium"),
            "options": question_doc.get("options", []),
            "images": merged_images,  # Include both option images and figures
            "questionFigures": figures_with_urls,  # Separate field for diagrams/figures
            "enhancedOptions": question_doc.get("enhancedOptions"),
            "correctAnswer": question_doc.get("correctAnswer") or question_doc.get("correct_answer"),  # Include answer for debugging
            "evaluationMode": evaluation_mode,
            "evaluation_mode": evaluation_mode,
            "metadata": question_doc.get("metadata", {})
        }

        # Format LaTeX expressions in question text and options
        question = format_question_latex(question)
        
        logger.info(f"Returning question {question_id}: {len(images_with_urls)} option images, {len(figures_with_urls)} figures")
        if figures_with_urls:
            for idx, fig in enumerate(figures_with_urls):
                logger.info(f"  Figure {idx + 1}: ID={fig.get('id')}, has_base64={bool(fig.get('base64Data'))}, base64_len={len(fig.get('base64Data', ''))}")
        
        return {
            "success": True,
            "question": question
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get next practice question error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get next question: {str(e)}"
        )


class UploadedImageFile(BaseModel):
    """Model for uploaded image files from frontend"""
    data: str  # Base64 encoded image data
    name: str
    type: str  # MIME type


class UploadedDocumentFile(BaseModel):
    """Model for uploaded document files (PDF/DOCX) from frontend"""
    data: str  # Base64 encoded document data
    name: str
    type: str  # 'pdf' or 'docx'


class QuestionPageRefsModel(BaseModel):
    """Per-question page mapping from the Stoody Pen QuestionSession."""
    activePages: Optional[List[int]] = None        # Physical notebook page numbers
    bookType: Optional[str] = None                 # e.g. "LS", "MS"
    copyId: Optional[str] = None                   # Copy set ID
    timeIntervals: Optional[List[Dict[str, Any]]] = None  # [{startTs, endTs}]


class EvaluateRequest(BaseModel):
    questionId: str
    answerText: Optional[str] = None
    canvasData: Optional[str] = None
    canvasPages: Optional[List[str]] = None
    # Uploaded files from mobile/desktop file picker
    uploadedImages: Optional[List[UploadedImageFile]] = None  # Images uploaded directly
    uploadedDocuments: Optional[List[UploadedDocumentFile]] = None  # PDF/DOCX files
    # Optional tracking fields for practice history
    documentId: Optional[str] = None  # Practice set document ID
    sessionId: Optional[str] = None   # Practice session ID
    timeSpent: Optional[int] = None   # Time spent in seconds
    hintsUsed: Optional[int] = 0      # Number of hints used
    # Per-question page mapping from Stoody Pen (which pages + time intervals)
    questionPageRefs: Optional[QuestionPageRefsModel] = None

    # Be flexible: accept pages as strings or objects with common keys; normalize to data URLs
    @validator('canvasData', pre=True)
    def _normalize_canvas_data(cls, v):
        try:
            if v and isinstance(v, str) and not v.startswith('data:image'):
                return f"data:image/png;base64,{v}"
        except Exception:
            pass
        return v

    @validator('canvasPages', pre=True)
    def _normalize_canvas_pages(cls, v):
        if v is None:
            return v
        try:
            if isinstance(v, list):
                out: List[str] = []
                for item in v:
                    s = None
                    if isinstance(item, str):
                        s = item
                    elif isinstance(item, dict):
                        s = (
                            item.get('dataUrl')
                            or item.get('url')
                            or item.get('data')
                            or item.get('image')
                            or item.get('src')
                        )
                    if s:
                        if not s.startswith('data:image'):
                            s = f"data:image/png;base64,{s}"
                        out.append(s)
                return out
            # If a single string is provided, wrap as list
            if isinstance(v, str):
                s = v
                if not s.startswith('data:image'):
                    s = f"data:image/png;base64,{s}"
                return [s]
        except Exception:
            return v
        return v

    @root_validator(pre=True)
    def _coerce_aliases(cls, values):
        # Accept snake_case aliases from frontend
        mapping = {
            'question_id': 'questionId',
            'answer_text': 'answerText',
            'canvas_data': 'canvasData',
            'canvas_pages': 'canvasPages',
            'document_id': 'documentId',
            'session_id': 'sessionId',
            'time_spent': 'timeSpent',
            'hints_used': 'hintsUsed',
            'question_page_refs': 'questionPageRefs',
        }
        for src, dst in mapping.items():
            if src in values and dst not in values:
                values[dst] = values[src]
        # If canvasPages provided as single object/string elsewhere, normalize to list
        cp = values.get('canvasPages')
        if isinstance(cp, str):
            values['canvasPages'] = [cp]
        return values

class EvaluateResponse(BaseModel):
    success: bool = True
    evaluation: Dict[str, Any]


def _build_evaluation_prompt(
    *,
    question_text: str,
    options_text: str,
    correct_answer: str,
    correct_answer_value: str,
    is_option_letter: bool,
    is_mcq: bool,
    answer_text: str,
    uploaded_doc_text: str,
    num_student_images: int,
    num_question_figures: int,
    num_option_images: int,
) -> str:
    """Build a single, universal evaluation prompt.

    Vision LLMs handle MCQ, paragraph, math derivation, and diagram answers without
    type-specific branching when the prompt is clean and the rules are uniform.
    """
    parts: List[str] = []
    num_q_images = num_question_figures + num_option_images
    total_images = num_q_images + num_student_images

    # Per-image labelling. Images are sent in this order to match the labels:
    #   1) question figures, 2) option images, 3) student pages.
    # Strict labelling stops the LLM from describing question content as if it were
    # student work (e.g. seeing a benzene ring in option B's image and reporting
    # "the student drew a benzene ring").
    if total_images > 0:
        guide: List[str] = ["IMAGES IN THIS REQUEST (read in order):"]
        idx = 1
        for fi in range(num_question_figures):
            guide.append(f"  Image {idx} — QUESTION FIGURE {fi + 1} (part of the question itself)")
            idx += 1
        for oi in range(num_option_images):
            opt_letter = chr(ord("A") + oi)
            guide.append(f"  Image {idx} — OPTION {opt_letter} (one of the multiple-choice options)")
            idx += 1
        for sp in range(num_student_images):
            guide.append(
                f"  Image {idx} — STUDENT'S ANSWER, page {sp + 1} of {num_student_images}"
            )
            idx += 1
        guide.append("")
        if num_q_images > 0 and num_student_images > 0:
            guide.append(
                "CRITICAL: question figures and option images are part of the question — "
                "they are NOT the student's work. Describe ONLY the student's pages in "
                "`work_shown`. Use the question images to understand what was asked, and "
                "use option images (when present) to compare against the student's answer."
            )
        parts.append("\n".join(guide))

    # Question text — when the question is presented entirely as images, say so explicitly
    # so the model reads the question from the figures instead of getting confused by an
    # empty QUESTION: block.
    cleaned_q_text = (question_text or "").strip()
    if cleaned_q_text:
        parts.append("\nQUESTION:")
        parts.append(question_text)
        if num_question_figures > 0:
            parts.append(
                "(The question text above is supplemented by the question figure(s) shown. "
                "Use both together to understand what is being asked.)"
            )
    elif num_question_figures > 0:
        parts.append("\nQUESTION:")
        parts.append(
            "(This question is presented entirely as an image. Read the question content "
            "from the QUESTION FIGURE image(s) listed above. There is no separate question text.)"
        )
    else:
        parts.append("\nQUESTION:")
        parts.append("(Question text not available — proceed with whatever context is present.)")

    if options_text:
        parts.append("\nOPTIONS:")
        parts.append(options_text)
        if num_option_images > 0:
            parts.append(
                "(The text above lists the options; the corresponding option images are also "
                "provided. When the student's answer is a drawing, compare it against the "
                "option images to identify which option they intended.)"
            )

    submission_lines: List[str] = []
    if answer_text:
        submission_lines.append(f"Typed answer: {answer_text}")
    if uploaded_doc_text:
        submission_lines.append(
            f"Uploaded document content:\n{_truncate_for_prompt(uploaded_doc_text, 8000)}"
        )
    if num_student_images:
        submission_lines.append(
            f"Handwritten canvas: {num_student_images} page(s) — read carefully across all pages."
        )
    if not submission_lines:
        submission_lines.append("(No answer submitted)")
    parts.append("\nSTUDENT'S SUBMISSION:")
    parts.append("\n".join(submission_lines))

    if correct_answer:
        parts.append("\nCORRECT ANSWER (reference for grading):")
        if is_mcq and is_option_letter:
            parts.append(f"Option {correct_answer}")
            if correct_answer_value and correct_answer_value != correct_answer:
                parts.append(f"Meaning of option {correct_answer}: {correct_answer_value}")
            parts.append(
                "Note: this is a multiple-choice question. The student's answer is correct "
                "if it matches option " + correct_answer + " in any form (the letter itself, "
                "the meaning above, the corresponding name/formula/value, or — when options "
                "are images — a drawing or structure that depicts the same thing as option " +
                correct_answer + "'s image)."
            )
        else:
            parts.append(correct_answer)
    else:
        parts.append(
            "\nNo direct teacher answer is available. The reference shown above is the "
            "question's cached model answer; treat it as the authoritative correct answer."
        )

    parts.append('''
HOW TO EVALUATE — read this carefully:

This is a SUBJECTIVE evaluation system. The student may express the correct answer
in any number of valid forms. Your job is in two steps:

  STEP 1 — IDENTIFY what the student communicated.
    Look at every page of their work. Describe concretely what is on the page:
    transcribe text and equations, describe drawings and diagrams in domain terms
    (e.g. "a 4-carbon skeletal formula with an –OH group on the second carbon",
    "a labeled animal cell showing nucleus, mitochondria, and Golgi apparatus",
    "the student integrated by parts: u = x, dv = e^x dx, then …"). Use the
    domain vocabulary that fits the question (chemistry, biology, math, physics,
    history, etc.). Put this description in `work_shown`.

  STEP 2 — JUDGE semantic equivalence to the correct answer.
    The student is CORRECT if what they communicated means the same thing as
    the correct answer, in ANY form:
      - The option letter (e.g. "D")
      - The text/value/name of the correct option (e.g. "2-butanol", "butan-2-ol",
        "sec-butanol", "CH₃CH(OH)CH₂CH₃")
      - A drawing, structure, or diagram that represents the correct answer
        (e.g. a skeletal formula showing the same molecule)
      - A worked solution that arrives at the correct answer
      - A paragraph that conveys the correct concept
      - A numerically equivalent value ("9" = "nine" = "9.0"; "1/2" = "0.5";
        "7 days" = "7" when the question is about days)

    Mark wrong only when there is a CLEAR meaning mismatch — different number,
    different molecule, different concept. If the student's work shows wrong
    reasoning but happens to land on a value that matches by coincidence,
    that's still wrong — explain why in `what_went_wrong`.

    If the handwriting is genuinely unreadable, do NOT guess. Set
    `extracted_answer` to "", `is_correct` to false, and in `what_went_wrong`
    state specifically what you could see, what was unclear, and what the
    student could do (e.g. "I could see a structural formula in the lower
    half of the page but the labels next to the carbon chain were too faint
    to read. Please rewrite the labels more clearly or write the option
    letter / IUPAC name of your answer.").

WORK_SHOWN field — important:
  This is shown to the student so they can confirm the system read their work.
  It must DESCRIBE concretely, NOT judge correctness:
    - Text/equations: transcribe what was written.
    - Drawings/structures/diagrams: describe in domain-appropriate detail
      (carbon count, functional groups, labels, organelles, vectors, etc.).
    - If student wrote across multiple pages, mention which page contained what.
    - Be specific enough that the student recognizes their own work.

EXTRACTED_ANSWER field:
  The student's FINAL answer in its most concise form. For an MCQ where the
  student drew the structure of option D, `extracted_answer` should be "D"
  (or the IUPAC name) — i.e. the answer reduced to a comparable form, not the
  whole drawing description (that goes in `work_shown`).

OUTPUT — strict JSON only (no markdown fences, no commentary, no text outside the JSON):
{
  "is_correct": true | false,
  "score": 0.0 to 1.0,
  "extracted_answer": "the student's final answer in its most concise comparable form",
  "work_shown": "concrete domain-appropriate description of what the student wrote/drew",
  "what_went_wrong": "if wrong: specific mistake the student made. If unreadable: what you could see + what was unclear + what student should do. Empty string if correct."
}''')

    parts.append("\nNotes for the JSON:")
    parts.append("  - is_correct must be a boolean (true/false), not a string.")
    parts.append("  - score must be a number 0.0–1.0; clamp it inside that range.")
    parts.append("  - Formatting: use $...$ for inline math or chemistry notation (e.g. $x_0$,")
    parts.append("    $\\frac{a}{b}$, $\\ce{H2O}$) and $$...$$ for display math. Do NOT use")
    parts.append("    \\(...\\) or \\[...\\]. Wrap formulas and identifiers in $...$ so the renderer")
    parts.append("    typesets them correctly.")
    parts.append("  - Use double quotes for all strings; escape internal quotes as \\\".")
    parts.append("  - Do NOT include `correct_solution`, `feedback`, or `reasoning` in the output —")
    parts.append("    those are handled separately and would just waste tokens here.")

    return "\n".join(parts)


def _build_case_study_evaluation_prompt(
    *,
    question_text: str,
    options_text: str,
    answer_text: str,
    uploaded_doc_text: str,
    num_student_images: int,
    num_question_figures: int,
    num_option_images: int,
    model_solution_text: str,
    teacher_reference_answer: str,
) -> str:
    """Build a case-study rubric prompt that grades structure and reasoning."""
    parts: List[str] = []
    num_q_images = num_question_figures + num_option_images
    total_images = num_q_images + num_student_images

    if total_images > 0:
        guide: List[str] = ["IMAGES IN THIS REQUEST (read in order):"]
        idx = 1
        for fi in range(num_question_figures):
            guide.append(f"  Image {idx} - QUESTION / CASE FIGURE {fi + 1}")
            idx += 1
        for oi in range(num_option_images):
            opt_letter = chr(ord("A") + oi)
            guide.append(f"  Image {idx} - OPTION {opt_letter} / CASE MATERIAL")
            idx += 1
        for sp in range(num_student_images):
            guide.append(
                f"  Image {idx} - STUDENT'S ANSWER, page {sp + 1} of {num_student_images}"
            )
            idx += 1
        guide.append("")
        guide.append(
            "Use question/case images only as context. In `work_shown`, describe only "
            "the student's submitted pages and uploaded answer material."
        )
        parts.append("\n".join(guide))

    parts.append("\nCASE QUESTION:")
    parts.append((question_text or "").strip() or "(Question text not available.)")
    if options_text:
        parts.append("\nOPTIONS / CASE MATERIAL:")
        parts.append(options_text)

    submission_lines: List[str] = []
    if answer_text:
        submission_lines.append(f"Typed answer: {answer_text}")
    if uploaded_doc_text:
        submission_lines.append(
            f"Uploaded document content:\n{_truncate_for_prompt(uploaded_doc_text, 8000)}"
        )
    if num_student_images:
        submission_lines.append(
            f"Handwritten canvas: {num_student_images} page(s) - read the response structure carefully."
        )
    if not submission_lines:
        submission_lines.append("(No answer submitted)")
    parts.append("\nSTUDENT'S SUBMISSION:")
    parts.append("\n".join(submission_lines))

    if teacher_reference_answer:
        parts.append("\nTEACHER REFERENCE / RUBRIC:")
        parts.append(teacher_reference_answer)
        parts.append(
            "Use this as guidance, not an exact wording match. Award credit for equivalent "
            "business reasoning, prioritization, and executable recommendations."
        )
    if model_solution_text:
        parts.append("\nMODEL MBA-STYLE ANSWER / RUBRIC:")
        parts.append(_truncate_for_prompt(model_solution_text, 7000))

    parts.append('''
HOW TO EVALUATE THIS CASE STUDY:

Grade the student's response like an MBA / business-school case answer, not like a
single-answer STEM problem. Score the quality of thinking, structure, and execution
logic. Use this rubric:

  1. Problem diagnosis (20%): identifies the core causes, not only symptoms.
  2. Growth levers (20%): proposes relevant acquisition, retention, pricing, product,
     channel, conversion, or operations levers.
  3. Prioritization and tradeoffs (15%): explains what to do first and why.
  4. Metrics (15%): names useful KPIs such as revenue growth, CAC, LTV, conversion,
     retention, repeat purchase, AOV, contribution margin, churn, payback period.
  5. 90-day execution plan (20%): gives a practical 30/60/90 day or phased action plan.
  6. Clarity and assumptions (10%): communicates cleanly and states risks/assumptions.

`is_correct` means the answer is acceptable / passing overall. Set it to true when
score >= 0.60, false when score < 0.60.

Format all string fields as concise Markdown without tables.

`extracted_answer` should be 1-2 short sentences summarizing the student's case response.
`work_shown` should be a readable bullet list of what the student actually covered, grouped
by themes when useful. Prefer this format:
- **Pros:** ...
- **Cons:** ...
- **Framework:** ...
- **TAM / Market:** ...
- **Business model:** ...
- **Competition:** ...

`what_went_wrong` should be a readable bullet list of the main gaps and improvements.
Use short, concrete bullets such as:
- **Diagnosis:** ...
- **Growth levers:** ...
- **Metrics:** ...
- **Execution plan:** ...
- **Risks / assumptions:** ...
If score >= 0.90 this may be an empty string; otherwise include concrete improvement advice.

OUTPUT - strict JSON only (no markdown fences, no commentary, no text outside JSON):
{
  "is_correct": true | false,
  "score": 0.0 to 1.0,
  "extracted_answer": "1-2 sentence overall assessment",
  "work_shown": "- **Theme:** concise point\\n- **Theme:** concise point",
  "what_went_wrong": "- **Gap:** concise improvement\\n- **Gap:** concise improvement"
}''')

    return "\n".join(parts)


def _build_evaluation_system_prompt(detected_language: str) -> str:
    """Subjective-first system prompt for the evaluator.

    Core stance: the student may express a correct answer in many forms (letter, value,
    name, formula, structural drawing, diagram, paragraph, derivation). The evaluator's
    job is to identify what they communicated, then judge semantic equivalence to the
    correct answer — not to search for a particular form like a letter.
    """
    lang_rule = ""
    if detected_language == "hindi":
        lang_rule = (
            "CRITICAL: The question is in Hindi (Devanagari script). Respond ENTIRELY in Hindi "
            "(work_shown and what_went_wrong fields must be in Hindi). "
        )
    return (
        f"{lang_rule}"
        "You are an expert tutor grading a student's handwritten answer. You will see the "
        "question (with any diagrams or option images) and the student's work across one or "
        "more pages. "
        "Approach this as a subjective evaluation: the student may express the correct answer "
        "in any valid form — a letter, a value, a name, a formula, a structural drawing, a "
        "labeled diagram, a worked solution, or a paragraph. First identify what the student "
        "actually communicated (describe drawings in domain terms, transcribe text). Then "
        "judge whether their communication is semantically equivalent to the correct answer. "
        "Accept any equivalent form. "
        "If the handwriting is genuinely unreadable, do NOT guess what the student wrote based "
        "on what the answer should be — that is confirmation bias. Instead say specifically "
        "what you could see and what was unclear. "
        "FORMATTING: use $...$ for inline math or chemistry notation (e.g. $x_0$, "
        "$\\frac{a}{b}$, $\\ce{H2O}$) and $$...$$ for display math. Never use \\(...\\) "
        "or \\[...\\]. Wrap formulas and identifiers in $...$ so the renderer typesets them "
        "correctly. Use LaTeX commands (\\frac, \\sqrt, \\alpha, \\ce), not raw Unicode "
        "for symbols that need typesetting. "
        "OUTPUT: only valid JSON with the five required keys (is_correct, score, "
        "extracted_answer, work_shown, what_went_wrong). No markdown fences, no commentary, "
        "no extra fields."
    )


def _build_case_study_evaluation_system_prompt(detected_language: str) -> str:
    lang_rule = ""
    if detected_language == "hindi":
        lang_rule = (
            "CRITICAL: The question is in Hindi (Devanagari script). Respond ENTIRELY in Hindi "
            "(all JSON string fields must be in Hindi). "
        )
    return (
        f"{lang_rule}"
        "You are an MBA case-study evaluator grading a student's handwritten or typed answer. "
        "Do not look for one exact answer. Judge the quality of diagnosis, business logic, "
        "prioritization, metrics, execution plan, assumptions, and clarity. "
        "Be fair to different valid strategies, but be strict about vague answers that only "
        "list generic actions without reasoning or metrics. "
        "OUTPUT: only valid JSON with the five required keys (is_correct, score, "
        "extracted_answer, work_shown, what_went_wrong). No markdown fences, no commentary, "
        "no extra fields."
    )


def _parse_evaluation_response(
    *,
    raw_response: str,
    correct_answer_display: str,
    has_correct_answer: bool,
    answer_text: str,
    evaluation_mode: str = EVALUATION_MODE_STANDARD,
) -> Dict[str, Any]:
    """Parse the LLM evaluator JSON output.

    The evaluator now returns only verdict-related fields. The cached step-by-step
    `correctSolution` is stitched in by the caller after this function returns; we
    do NOT read `correct_solution`/`feedback`/`reasoning`/`solved_answer` from the
    LLM output (those keys are no longer requested in the prompt).
    """
    evaluation_data: Dict[str, Any] = {
        "correct": False,
        "score": 0.0,
        "extractedAnswer": "",
        "workShown": "",
        "whatWentWrong": "",
        "correctSolution": "",  # filled in by caller from solution_cache
        "correctAnswer": correct_answer_display if has_correct_answer else "",
        "correctAnswerSource": "admin_provided" if has_correct_answer else "unknown",
        "evaluationMode": evaluation_mode,
        "evaluation_mode": evaluation_mode,
    }

    parsed = robust_json_parse(raw_response) if raw_response else None

    if not parsed or not isinstance(parsed, dict):
        logger.warning(f"⚠️ Could not parse JSON from LLM. Raw (first 200): {raw_response[:200]!r}")
        evaluation_data["whatWentWrong"] = (
            "We could not read the submission clearly. Please try writing again."
            if evaluation_mode != EVALUATION_MODE_CASE_STUDY
            else "We could not read the case-study response clearly. Please rewrite it with clear sections."
        )
        evaluation_data["extractedAnswer"] = answer_text or ""
        return evaluation_data

    is_correct_val = parsed.get("is_correct", parsed.get("correct", False))
    if isinstance(is_correct_val, bool):
        evaluation_data["correct"] = is_correct_val
    elif isinstance(is_correct_val, str):
        evaluation_data["correct"] = is_correct_val.strip().lower() in ("true", "yes", "correct", "1")
    else:
        evaluation_data["correct"] = bool(is_correct_val)

    score_val = parsed.get("score")
    if score_val is None:
        evaluation_data["score"] = 1.0 if evaluation_data["correct"] else 0.0
    else:
        try:
            evaluation_data["score"] = float(score_val)
        except (TypeError, ValueError):
            evaluation_data["score"] = 1.0 if evaluation_data["correct"] else 0.0
    evaluation_data["score"] = max(0.0, min(1.0, evaluation_data["score"]))

    evaluation_data["extractedAnswer"] = _normalize_latex_for_render(
        str(parsed.get("extracted_answer", "")).strip() or (answer_text or "")
    )
    evaluation_data["workShown"] = _normalize_latex_for_render(
        str(parsed.get("work_shown", "")).strip()
    )
    evaluation_data["whatWentWrong"] = _normalize_latex_for_render(
        str(parsed.get("what_went_wrong", "")).strip()
    )

    return evaluation_data


@router.post("/evaluate", response_model=EvaluateResponse)
@limiter.limit("120/minute")
async def evaluate_submission(
    request: Request,
    payload: EvaluateRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    """Evaluate a student's handwritten / typed submission via a single vision LLM call.

    Pipeline:
      1. Load question (text + figures + option images).
      2. Filter empty student canvas pages and apply a single light enhancement pass.
      3. Extract text from any uploaded PDFs/DOCX.
      4. Build ONE prompt covering MCQ, numerical, paragraph, math, and diagram answers.
      5. ONE LLM vision call (or text-only call when no images are involved).
      6. Parse JSON, save attempt to history, return.

    Returns: { success, evaluation: { correct, score, extractedAnswer, workShown,
              whatWentWrong, correctSolution, feedback, reasoning, correctAnswer, ... } }
    """
    try:
        qid = payload.questionId
        answer_text = (payload.answerText or "").strip()
        canvas_data = payload.canvasData

        # B2C detection (separate database)
        user_type = current_user.get("user_type", "")
        is_b2c = current_user.get("is_b2c", False) or user_type == "b2c_user"

        logger.info(f"📝 Evaluating Q:{qid} (user_type={user_type}, b2c={is_b2c})")

        # Normalize legacy single-canvas payload to a data URL
        if canvas_data and not canvas_data.startswith("data:image"):
            canvas_data = f"data:image/png;base64,{canvas_data}"

        # 1. Load the question
        question_doc = await _load_question_doc(db, qid, is_b2c=is_b2c)
        if not question_doc:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Question not found")

        ca_primary = question_doc.get("correctAnswer")
        ca_alt = question_doc.get("correct_answer")
        correct_answer = str(
            ca_primary if ca_primary is not None else (ca_alt if ca_alt is not None else "")
        ).strip()

        # Resolve option letter to actual value (e.g. "A" -> "7 days")
        resolved = _resolve_correct_answer(correct_answer, question_doc)
        correct_answer_value = resolved["resolved_value"]   # The actual answer content
        correct_answer_display = resolved["display"]         # Human-readable: "A (7 days)"
        is_option_letter = resolved["is_option_letter"]      # True if answer is A/B/C/D
        correct_answer_value = _normalize_latex_for_render(str(correct_answer_value or "").strip())
        correct_answer_display = _normalize_latex_for_render(str(correct_answer_display or "").strip())

        # Extract question text and options
        question_text = str(question_doc.get("text", ""))
        options_text = _options_text_from_question(question_doc)
        
        # Check stored question_type (set during upload or admin edit)
        stored_question_type = (question_doc.get("question_type") or "").lower().strip()

        # Determine if this is MCQ: must have options AND not be explicitly marked subjective
        is_mcq = bool(options_text) and stored_question_type != "subjective"

        # For subjective/non-MCQ questions: if the stored correctAnswer is a single option
        # letter (A-J) but there are no options to resolve it against, it's a meaningless
        # OCR artifact. Clear it so the LLM solves the question itself.
        if not is_mcq and is_option_letter and correct_answer_value == correct_answer:
            logger.info(
                f"Clearing bogus option-letter correctAnswer='{correct_answer}' for non-MCQ Q:{qid} "
                f"(no options to resolve against — LLM will solve instead)."
            )
            correct_answer = ""
            correct_answer_value = ""
            correct_answer_display = ""
            is_option_letter = False

        evaluation_mode = _infer_evaluation_mode(question_doc, is_mcq=is_mcq)
        logger.info(f"Evaluation mode for Q:{qid}: {evaluation_mode}")

        # 2. Load question images (figures + option images)
        question_figure_images = await _figure_images_base64(question_doc, db, is_b2c)
        option_images_data = await _option_images_base64(question_doc, db, is_b2c)
        option_images = [oi["image"] for oi in option_images_data]
        all_question_images = question_figure_images + option_images

        # 3. Collect student images.
        # IMPORTANT distinction:
        #   - canvasPages: rendered from pen strokes — already pixel-perfect crisp lines on
        #     a clean white background. These must NOT go through contrast/sharpen because
        #     those filters introduce edge artifacts on already-sharp thin strokes (e.g. a
        #     hand-drawn "O" can read as "D" after Sharpen widens its loop).
        #   - uploadedImages: real photos of paper. Lighting and focus may be poor, so a
        #     light enhancement pass genuinely helps the model read them.
        # We process the two paths separately so each gets the treatment it actually needs.
        canvas_pages_raw: List[str] = []
        if payload.canvasPages:
            canvas_pages_raw = list(payload.canvasPages)
        elif canvas_data:
            canvas_pages_raw = [canvas_data]

        uploaded_photos_raw: List[str] = []
        if payload.uploadedImages:
            for uploaded_img in payload.uploadedImages:
                data = uploaded_img.data
                if not data.startswith("data:"):
                    body = data.split(",")[-1] if "," in data else data
                    data = f"data:{uploaded_img.type};base64,{body}"
                uploaded_photos_raw.append(data)

        # 4. Filter empty canvas pages; do not enhance them (they're already crisp).
        student_images: List[str] = []
        try:
            from utils.image_processor import enhance_canvas_images_batch, is_canvas_empty
        except ImportError as ie:
            logger.warning(f"Image processor not available: {ie}; using raw images.")
            enhance_canvas_images_batch = None  # type: ignore
            is_canvas_empty = None  # type: ignore

        if canvas_pages_raw:
            if is_canvas_empty:
                non_empty = [img for img in canvas_pages_raw if img and not is_canvas_empty(img)]
                if not non_empty:
                    # All blank — keep raw so the LLM can correctly say "no answer".
                    non_empty = [img for img in canvas_pages_raw if img]
            else:
                non_empty = [img for img in canvas_pages_raw if img]
            student_images.extend(non_empty)

        # Uploaded photos: light enhancement helps (camera-quality variability).
        if uploaded_photos_raw:
            if enhance_canvas_images_batch:
                try:
                    enhanced_photos = (
                        enhance_canvas_images_batch(uploaded_photos_raw, target_width=1500)
                        or uploaded_photos_raw
                    )
                except Exception as enhance_err:
                    logger.warning(f"Photo enhancement failed: {enhance_err}; using raw images.")
                    enhanced_photos = uploaded_photos_raw
            else:
                enhanced_photos = uploaded_photos_raw
            student_images.extend(enhanced_photos)

        # 5. Extract text from any uploaded PDFs/DOCX
        uploaded_doc_text = ""
        if payload.uploadedDocuments:
            for doc in payload.uploadedDocuments:
                try:
                    text = await _extract_text_from_document(doc.data, doc.type, doc.name)
                    if text:
                        uploaded_doc_text += f"\n[From {doc.name}]:\n{text}\n"
                except Exception as doc_err:
                    logger.warning(f"Failed to extract text from {doc.name}: {doc_err}")

        has_correct_answer = bool(correct_answer)
        detected_language = detect_language(question_text)
        num_student_images = len(student_images)
        num_question_figures = len(question_figure_images)
        num_option_images = len(option_images)

        # Diagnostic: image dimensions per page so we can confirm the frontend is shipping
        # tightly-cropped student images (button-tap suppression + bbox crop both working).
        # Pair this with the [renderStrokeElementsToImage] log on the frontend; together they
        # tell you how many strokes were dropped by the export safety net and what size
        # image actually reached the LLM.
        student_image_sizes: List[str] = []
        for img in student_images:
            try:
                # Quick dimension probe via base64 length is too coarse; use PIL only when
                # we already have the image processor imported. Cheap fallback: just record
                # the data-URL length as a proxy for size.
                length_kb = round(len(img) / 1024) if img else 0
                student_image_sizes.append(f"{length_kb}kb")
            except Exception:
                student_image_sizes.append("?")

        logger.info(
            f"📷 Q:{qid} images — student={num_student_images} {student_image_sizes}, "
            f"question_figures={num_question_figures}, option_images={num_option_images}, "
            f"is_mcq={is_mcq}, mode={evaluation_mode}, has_ref_answer={has_correct_answer}, "
            f"lang={detected_language}"
        )

        # 5b. Resolve the cached step-by-step solution. This is shared across all students
        # who attempt the same question; the helper hits the LLM only on cache miss / hash
        # mismatch. The student's canvas pages are intentionally NOT passed in so a wrong
        # attempt cannot pollute the shared solution.
        cached = await _get_or_generate_solution(
            question_doc, db, current_user, is_b2c, evaluation_mode=evaluation_mode
        )
        cached_solution_text = cached.get("solution") or ""
        cached_final_answer = (cached.get("final_answer") or "").strip()

        # If admin didn't provide a correctAnswer but the cache has a final_answer, use it
        # as the evaluator's reference. The full step-by-step is shown to the student
        # post-evaluation; the evaluator only needs the short reference for the verdict.
        if evaluation_mode != EVALUATION_MODE_CASE_STUDY and not has_correct_answer and cached_final_answer:
            correct_answer = cached_final_answer
            correct_answer_value = _normalize_latex_for_render(cached_final_answer)
            correct_answer_display = _normalize_latex_for_render(cached_final_answer)
            has_correct_answer = True

        # 6. Build the single, unified prompt + system prompt
        if evaluation_mode == EVALUATION_MODE_CASE_STUDY:
            prompt = _build_case_study_evaluation_prompt(
                question_text=question_text,
                options_text=options_text,
                answer_text=answer_text,
                uploaded_doc_text=uploaded_doc_text,
                num_student_images=num_student_images,
                num_question_figures=num_question_figures,
                num_option_images=num_option_images,
                model_solution_text=cached_solution_text,
                teacher_reference_answer=correct_answer_display or correct_answer_value or correct_answer,
            )
            system_prompt = _build_case_study_evaluation_system_prompt(detected_language)
        else:
            prompt = _build_evaluation_prompt(
                question_text=question_text,
                options_text=options_text,
                correct_answer=correct_answer,
                correct_answer_value=correct_answer_value,
                is_option_letter=is_option_letter,
                is_mcq=is_mcq,
                answer_text=answer_text,
                uploaded_doc_text=uploaded_doc_text,
                num_student_images=num_student_images,
                num_question_figures=num_question_figures,
                num_option_images=num_option_images,
            )
            system_prompt = _build_evaluation_system_prompt(detected_language)

        # 7. ONE LLM call. Output is now small (5 fields, no solution/feedback/reasoning),
        # so 1000 tokens is plenty.
        # Image order MUST match the per-image labelling produced by the prompt builder:
        #   question figures first, then option images, then student pages (one per page).
        all_images = all_question_images + student_images
        if all_images:
            response = await _gate_vision_call(
                db, current_user, all_images,
                prompt,
                system_prompt=system_prompt,
                max_tokens=1000,
                temperature=0.2,
            )
        else:
            response = await _gate_text_call(
                db, current_user, prompt,
                system_prompt=system_prompt,
                max_tokens=1000,
                temperature=0.2,
            )

        raw_response = (response.get("response") or "").strip()
        logger.info(f"📥 LLM response (first 300): {raw_response[:300]!r}")

        # 8. Parse the verdict + work_shown + what_went_wrong, then stitch in the cached solution.
        evaluation_data = _parse_evaluation_response(
            raw_response=raw_response,
            correct_answer_display=correct_answer_display,
            has_correct_answer=has_correct_answer,
            answer_text=answer_text,
            evaluation_mode=evaluation_mode,
        )
        evaluation_data["correctSolution"] = cached_solution_text
        if cached.get("source") and not has_correct_answer:
            evaluation_data["correctAnswerSource"] = cached["source"]

        logger.info(
            f"🎯 Evaluation Q:{qid}: correct={evaluation_data['correct']}, "
            f"score={evaluation_data['score']:.2f}, "
            f"extracted={(evaluation_data.get('extractedAnswer') or '')[:60]!r}"
        )

        # === SAVE PRACTICE ATTEMPT TO DATABASE FOR HISTORY ===
        try:
            user_id = current_user.get("user_id") or current_user.get("student_id") or current_user.get("id")
            
            # Get document_id from question metadata or request payload
            doc_id = payload.documentId
            if not doc_id:
                # Try to extract from question metadata
                meta = question_doc.get("metadata", {}) or {}
                doc_id = meta.get("document_id") or meta.get("documentId") or question_doc.get("document_id")
            
            # Prepare attempt record
            q_type = "case_study" if evaluation_mode == EVALUATION_MODE_CASE_STUDY else ("mcq" if is_mcq else "subjective")

            # Build question_page_refs from the Stoody Pen QuestionSession
            question_page_refs = None
            if payload.questionPageRefs:
                qpr = payload.questionPageRefs
                if qpr.activePages:
                    question_page_refs = {
                        "active_pages": qpr.activePages,
                        "book_type": qpr.bookType,
                        "copy_id": qpr.copyId,
                        "time_intervals": qpr.timeIntervals or [],
                    }

            practice_attempt = {
                "student_id": str(user_id),
                "session_id": payload.sessionId,
                "document_id": doc_id,
                "question_id": qid,
                "question_text": question_text[:2000] if question_text else "",
                "question_type": q_type,
                "evaluation_mode": evaluation_mode,
                "options": question_doc.get("options"),
                "student_answer": evaluation_data.get("extractedAnswer", ""),
                "correct_answer": correct_answer,
                "is_correct": evaluation_data.get("correct", False),
                "score": evaluation_data.get("score", 0.0),
                "time_spent": payload.timeSpent,
                "hints_used": payload.hintsUsed or 0,
                "work_shown": evaluation_data.get("workShown", ""),
                "what_went_wrong": evaluation_data.get("whatWentWrong", ""),
                "correct_solution": evaluation_data.get("correctSolution", ""),
                "question_page_refs": question_page_refs,
                "created_at": datetime.utcnow(),
                "subject": question_doc.get("subject", ""),
                "difficulty": question_doc.get("difficulty", ""),
                "topic": (question_doc.get("metadata", {}) or {}).get("topic", ""),
            }
            
            # Save to appropriate database (B2C or main)
            if is_b2c:
                await db.b2c_insert_one("practice_attempts", practice_attempt)
            else:
                await db.mongo_insert_one("practice_attempts", practice_attempt)
            
            logger.info(f"💾 Saved practice attempt for student {user_id}, question {qid}")
            
        except Exception as save_err:
            # Log but don't fail the request if saving fails
            logger.error(f"❌ Failed to save practice attempt: {save_err}", exc_info=True)

        return EvaluateResponse(success=True, evaluation=evaluation_data)

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to evaluate submission: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to evaluate submission"
        )

@router.post("/sessions", response_model=SessionResponse)
@limiter.limit("30/minute")
async def start_practice_session(
    request: Request,
    session_data: StartSessionRequest,
    current_user: Dict[str, Any] = Depends(require_student_or_admin),
    db: DatabaseManager = Depends(get_database)
):
    """Start a new practice session"""
    try:
        user_id = current_user["user_id"]
        
        # Detect if user is B2C (uses B2C database)
        user_type = current_user.get("user_type", "")
        is_b2c = current_user.get("is_b2c", False) or user_type == "b2c_user"

        # Create session record
        session_record = {
            "student_id": user_id,
            "mode": session_data.mode,
            "subject": session_data.subject,
            "difficulty": session_data.difficulty,
            "time_limit": session_data.time_limit,
            "document_id": session_data.document_id,  # Track which practice set
            "questions_attempted": 0,
            "correct_answers": 0,
            "total_time_spent": 0,
            "started_at": datetime.utcnow(),
            "is_completed": False,
            "questions": []  # Will store question attempts
        }

        # Use B2C database for B2C users
        if is_b2c:
            session_id = await db.b2c_insert_one("practice_sessions", session_record)
        else:
            session_id = await db.mongo_insert_one("practice_sessions", session_record)

        if not session_id:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to start practice session"
            )

        return SessionResponse(
            id=session_id,
            mode=session_data.mode,
            subject=session_data.subject,
            difficulty=session_data.difficulty,
            questions_attempted=0,
            correct_answers=0,
            accuracy_rate=0.0,
            total_time_spent=0,
            started_at=session_record["started_at"],
            is_completed=False
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Start practice session error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to start practice session"
        )

@router.post("/sessions/{session_id}/answer")
@limiter.limit("200/minute")
async def submit_session_answer(
    request: Request,
    session_id: str,
    answer_data: SessionAnswer,
    current_user: Dict[str, Any] = Depends(require_student_or_admin),
    db: DatabaseManager = Depends(get_database)
):
    """Submit answer for a question in a practice session"""
    try:
        user_id = current_user["user_id"]

        # Get session
        session = await db.mongo_find_one("practice_sessions", {"_id": session_id})

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Practice session not found"
            )

        # Check ownership (students can only access their own sessions)
        if (current_user["user_type"] == "student" and
            session["student_id"] != user_id):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )

        if session["is_completed"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Session is already completed"
            )

        # Get question to validate answer
        question = await db.mongo_find_one("questions", {"question_id": answer_data.question_id})

        # Validate answer
        is_correct = False
        score = 0
        if question:
            correct_answer = question.get("correct_answer", "")
            is_correct = (answer_data.answer.strip().lower() == correct_answer.strip().lower())
            if is_correct:
                score = question.get("points", 4.0)  # Default 4 marks per question

        # Create question attempt record in session
        question_attempt = {
            "question_id": answer_data.question_id,
            "answer": answer_data.answer,
            "is_correct": is_correct,
            "time_spent": answer_data.time_spent,
            "answered_at": datetime.utcnow()
        }

        # Update session with new attempt
        update_data = {
            "$push": {"questions": question_attempt},
            "$inc": {
                "questions_attempted": 1,
                "total_time_spent": answer_data.time_spent
            }
        }

        if is_correct:
            update_data["$inc"]["correct_answers"] = 1

        await db.mongo_update_one(
            "practice_sessions",
            {"_id": session_id},
            update_data
        )

        # Track in question_attempts collection for student monitoring
        if current_user["user_type"] == "student":
            try:
                student_oid = ObjectId(user_id)

                # Get admin_id from JWT token for data isolation
                admin_id = current_user.get("admin_id")
                if not admin_id:
                    logger.warning(f"Student {user_id} has no admin_id in JWT token")
                    admin_id = None

                # Insert into question_attempts collection
                attempt_doc = {
                    "student_id": student_oid,
                    "question_id": answer_data.question_id,
                    "session_id": session_id,
                    "answer": answer_data.answer,
                    "is_correct": is_correct,
                    "score": score,
                    "time_spent": answer_data.time_spent,
                    "created_at": datetime.utcnow(),
                    "metadata": {
                        "subject": question.get("subject") if question else None,
                        "difficulty": question.get("difficulty") if question else None
                    }
                }

                # Add admin_id for data isolation if available
                if admin_id:
                    attempt_doc["admin_id"] = admin_id

                await db.mongo_insert_one("question_attempts", attempt_doc)

                # Log activity in student_activity_log
                activity_doc = {
                    "student_id": student_oid,
                    "action": "question_attempted",
                    "timestamp": datetime.utcnow(),
                    "metadata": {
                        "question_id": answer_data.question_id,
                        "session_id": session_id,
                        "is_correct": is_correct,
                        "score": score,
                        "time_spent": answer_data.time_spent
                    }
                }

                # Add admin_id for data isolation if available
                if admin_id:
                    activity_doc["admin_id"] = admin_id

                await db.mongo_insert_one("student_activity_log", activity_doc)
            except Exception as e:
                logger.warning(f"Failed to track question attempt: {str(e)}")

        return {
            "message": "Answer submitted successfully",
            "is_correct": is_correct,
            "question_id": answer_data.question_id,
            "score": score
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Submit session answer error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to submit answer"
        )

@router.post("/sessions/{session_id}/complete")
@limiter.limit("30/minute")
async def complete_practice_session(
    request: Request,
    session_id: str,
    current_user: Dict[str, Any] = Depends(require_student_or_admin),
    db: DatabaseManager = Depends(get_database)
):
    """Complete a practice session"""
    try:
        user_id = current_user["user_id"]

        # Get session
        session = await db.mongo_find_one("practice_sessions", {"_id": session_id})

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Practice session not found"
            )

        # Check ownership
        if (current_user["user_type"] == "student" and
            session["student_id"] != user_id):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )

        if session["is_completed"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Session is already completed"
            )

        # Mark session as completed
        await db.mongo_update_one(
            "practice_sessions",
            {"_id": session_id},
            {
                "$set": {
                    "is_completed": True,
                    "completed_at": datetime.utcnow()
                }
            }
        )

        # Get updated session
        updated_session = await db.mongo_find_one("practice_sessions", {"_id": session_id})

        accuracy_rate = 0.0
        if updated_session["questions_attempted"] > 0:
            accuracy_rate = (updated_session["correct_answers"] / updated_session["questions_attempted"]) * 100

        return SessionResponse(
            id=session_id,
            mode=updated_session["mode"],
            subject=updated_session.get("subject"),
            difficulty=updated_session.get("difficulty"),
            questions_attempted=updated_session["questions_attempted"],
            correct_answers=updated_session["correct_answers"],
            accuracy_rate=round(accuracy_rate, 1),
            total_time_spent=updated_session["total_time_spent"],
            started_at=updated_session["started_at"],
            completed_at=updated_session["completed_at"],
            is_completed=True
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Complete practice session error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to complete practice session"
        )

@router.get("/sessions", response_model=SessionsListResponse)
@limiter.limit("60/minute")
async def get_practice_sessions(
    request: Request,
    page: int = Query(1, ge=1),
    limit: int = Query(20, ge=1, le=100),
    mode: Optional[str] = Query(None),
    is_completed: Optional[bool] = Query(None),
    current_user: Dict[str, Any] = Depends(require_student_or_admin),
    db: DatabaseManager = Depends(get_database)
):
    """Get practice sessions"""
    try:
        user_id = current_user["user_id"]
        user_type = current_user["user_type"]

        # Build filter
        filter_dict = {}
        if user_type == "student":
            filter_dict["student_id"] = user_id

        if mode:
            filter_dict["mode"] = mode
        if is_completed is not None:
            filter_dict["is_completed"] = is_completed

        # Get total count
        all_sessions = await db.mongo_find("practice_sessions", filter_dict)
        total_sessions = len(all_sessions)

        # Get paginated results
        skip = (page - 1) * limit
        sessions_data = await db.mongo_find(
            "practice_sessions",
            filter_dict,
            sort=[("started_at", -1)],
            skip=skip,
            limit=limit
        )

        sessions = []
        for session in sessions_data:
            accuracy_rate = 0.0
            if session["questions_attempted"] > 0:
                accuracy_rate = (session["correct_answers"] / session["questions_attempted"]) * 100

            sessions.append(SessionResponse(
                id=str(session["_id"]),
                mode=session["mode"],
                subject=session.get("subject"),
                difficulty=session.get("difficulty"),
                questions_attempted=session["questions_attempted"],
                correct_answers=session["correct_answers"],
                accuracy_rate=round(accuracy_rate, 1),
                total_time_spent=session["total_time_spent"],
                started_at=session["started_at"],
                completed_at=session.get("completed_at"),
                is_completed=session["is_completed"]
            ))

        return SessionsListResponse(
            sessions=sessions,
            total=total_sessions,
            page=page,
            limit=limit
        )

    except Exception as e:
        logger.error(f"Get practice sessions error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to get practice sessions"
        )

@router.get("/stats", response_model=PracticeStats)
@limiter.limit("30/minute")
async def get_practice_stats(
    request: Request,
    current_user: Dict[str, Any] = Depends(require_student_or_admin),
    db: DatabaseManager = Depends(get_database),
    cache: CacheManager = Depends(get_cache)
):
    """Get practice statistics"""
    try:
        user_id = current_user["user_id"]
        user_type = current_user["user_type"]

        # Check cache first
        cache_key = f"practice_stats:{user_id}" if user_type == "student" else "practice_stats:admin"
        cached_stats = await cache.get(cache_key, "practice")
        if cached_stats:
            return PracticeStats(**cached_stats)

        # Build filter
        filter_dict = {}
        if user_type == "student":
            filter_dict["student_id"] = user_id

        # Get all sessions
        all_sessions = await db.mongo_find("practice_sessions", filter_dict)

        total_sessions = len(all_sessions)
        total_time_spent = sum(s.get("total_time_spent", 0) for s in all_sessions)

        # Calculate average accuracy
        completed_sessions = [s for s in all_sessions if s.get("is_completed", False)]
        total_accuracy = 0
        if completed_sessions:
            for session in completed_sessions:
                if session["questions_attempted"] > 0:
                    accuracy = (session["correct_answers"] / session["questions_attempted"]) * 100
                    total_accuracy += accuracy
            average_accuracy = total_accuracy / len(completed_sessions)
        else:
            average_accuracy = 0.0

        # Sessions by mode
        sessions_by_mode = {}
        for session in all_sessions:
            mode = session.get("mode", "unknown")
            sessions_by_mode[mode] = sessions_by_mode.get(mode, 0) + 1

        # Recent activity (last 7 days)
        recent_cutoff = datetime.utcnow() - timedelta(days=7)
        recent_sessions = [s for s in all_sessions if s["started_at"] >= recent_cutoff]
        recent_activity = [
            {
                "date": session["started_at"].date().isoformat(),
                "mode": session["mode"],
                "questions_attempted": session["questions_attempted"],
                "accuracy": round((session["correct_answers"] / session["questions_attempted"]) * 100, 1) if session["questions_attempted"] > 0 else 0
            }
            for session in recent_sessions[-10:]  # Last 10 recent sessions
        ]

        stats_data = {
            "total_sessions": total_sessions,
            "total_time_spent": total_time_spent,
            "average_accuracy": round(average_accuracy, 1),
            "sessions_by_mode": sessions_by_mode,
            "recent_activity": recent_activity
        }

        # Cache for 10 minutes
        await cache.set(cache_key, stats_data, 600, "practice")

        return PracticeStats(**stats_data)

    except Exception as e:
        logger.error(f"Practice stats error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to get practice statistics"
        )


@router.post("/grade", response_model=EvaluateResponse)
@limiter.limit("120/minute")
async def grade_submission(
    request: Request,
    payload: EvaluateRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    """Backward-compatible alias for the canonical /evaluate pipeline."""
    return await evaluate_submission(
        request=request,
        payload=payload,
        current_user=current_user,
        db=db
    )

# Backward compatibility endpoint for older clients that used legacy evaluator wiring.
@router.post("/evaluate-compat", response_model=EvaluateResponse)
@limiter.limit("120/minute")
async def evaluate_submission_compat(
    request: Request,
    payload: EvaluateRequest,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    return await grade_submission(request, payload, current_user, db)


# =============================================================================
# PRACTICE ATTEMPT HISTORY ENDPOINTS
# =============================================================================

class PracticeAttemptResponse(BaseModel):
    """Single practice attempt record"""
    id: str
    student_id: str
    question_id: str
    question_text: Optional[str] = None
    question_type: Optional[str] = None
    options: Optional[List[str]] = None
    student_answer: Optional[str] = None
    correct_answer: Optional[str] = None
    is_correct: bool = False
    score: float = 0.0
    time_spent: Optional[int] = None
    hints_used: int = 0
    evaluation_feedback: Optional[str] = None
    subject: Optional[str] = None
    difficulty: Optional[str] = None
    created_at: datetime
    document_id: Optional[str] = None

class PracticeHistoryStats(BaseModel):
    """Aggregated statistics for practice attempts"""
    total_attempted: int = 0
    total_correct: int = 0
    accuracy_percentage: float = 0.0
    avg_time_per_question: Optional[float] = None
    total_time_spent: int = 0

class PracticeHistoryResponse(BaseModel):
    """Response for practice attempt history"""
    success: bool = True
    attempts: List[Dict[str, Any]]
    total: int
    stats: PracticeHistoryStats

class PracticeSetStatsResponse(BaseModel):
    """Response for practice set statistics"""
    success: bool = True
    document_id: str
    question_stats: List[Dict[str, Any]]
    summary: Dict[str, Any]


@router.get("/attempts")
@limiter.limit("60/minute")
async def get_practice_attempts(
    request: Request,
    document_id: Optional[str] = Query(None, description="Filter by practice set document ID"),
    question_id: Optional[str] = Query(None, description="Filter by specific question ID"),
    subject: Optional[str] = Query(None, description="Filter by subject"),
    limit: int = Query(50, le=200, description="Maximum number of attempts to return"),
    offset: int = Query(0, ge=0, description="Offset for pagination"),
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    """
    Get practice attempt history for the current student.
    
    Returns a list of past practice attempts with:
    - Question details
    - Student's answer
    - Correct/incorrect status
    - Evaluation feedback
    - Aggregated statistics
    """
    try:
        user_id = current_user.get("user_id") or current_user.get("student_id") or current_user.get("id")
        
        # Detect if user is B2C
        user_type = current_user.get("user_type", "")
        is_b2c = current_user.get("is_b2c", False) or user_type == "b2c_user"
        
        # Build filter
        filter_dict = {"student_id": str(user_id)}
        if document_id:
            filter_dict["document_id"] = document_id
        if question_id:
            filter_dict["question_id"] = question_id
        if subject:
            filter_dict["subject"] = subject
        
        # Fetch attempts with pagination
        if is_b2c:
            attempts = await db.b2c_find(
                "practice_attempts",
                filter_dict,
                sort=[("created_at", -1)],
                limit=limit,
                skip=offset
            )
            total = await db.b2c_count("practice_attempts", filter_dict)
        else:
            attempts = await db.mongo_find(
                "practice_attempts",
                filter_dict,
                sort=[("created_at", -1)],
                limit=limit,
                skip=offset
            )
            total = await db.mongo_count("practice_attempts", filter_dict)
        
        # Convert ObjectId to string for JSON serialization
        attempt_list = []
        for a in attempts:
            attempt_dict = dict(a)
            if "_id" in attempt_dict:
                attempt_dict["id"] = str(attempt_dict.pop("_id"))
            if "created_at" in attempt_dict and attempt_dict["created_at"]:
                attempt_dict["created_at"] = attempt_dict["created_at"].isoformat()
            attempt_list.append(attempt_dict)
        
        # Calculate stats
        total_correct = sum(1 for a in attempt_list if a.get("is_correct"))
        total_time = sum(a.get("time_spent", 0) or 0 for a in attempt_list)
        
        stats = PracticeHistoryStats(
            total_attempted=total,
            total_correct=total_correct,
            accuracy_percentage=(total_correct / total * 100) if total > 0 else 0.0,
            avg_time_per_question=(total_time / len(attempt_list)) if attempt_list else None,
            total_time_spent=total_time
        )
        
        return {
            "success": True,
            "attempts": attempt_list,
            "total": total,
            "stats": stats.dict()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get practice attempts error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch practice attempts"
        )


@router.get("/attempts/{attempt_id}")
@limiter.limit("60/minute")
async def get_practice_attempt_detail(
    request: Request,
    attempt_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    """
    Get detailed information about a specific practice attempt.
    
    Returns the full attempt record including:
    - Original question with all details
    - Student's answer
    - Correct answer
    - Full AI feedback and reasoning
    - Work shown, what went wrong, correct solution
    """
    try:
        user_id = current_user.get("user_id") or current_user.get("student_id") or current_user.get("id")
        
        # Detect if user is B2C
        user_type = current_user.get("user_type", "")
        is_b2c = current_user.get("is_b2c", False) or user_type == "b2c_user"
        
        # Fetch the attempt
        from bson import ObjectId
        try:
            oid = ObjectId(attempt_id)
            filter_dict = {"_id": oid, "student_id": str(user_id)}
        except Exception:
            filter_dict = {"_id": attempt_id, "student_id": str(user_id)}
        
        if is_b2c:
            attempt = await db.b2c_find_one("practice_attempts", filter_dict)
        else:
            attempt = await db.mongo_find_one("practice_attempts", filter_dict)
        
        if not attempt:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Practice attempt not found"
            )
        
        # Convert for JSON serialization
        attempt_dict = dict(attempt)
        if "_id" in attempt_dict:
            attempt_dict["id"] = str(attempt_dict.pop("_id"))
        if "created_at" in attempt_dict and attempt_dict["created_at"]:
            attempt_dict["created_at"] = attempt_dict["created_at"].isoformat()
        
        return {
            "success": True,
            "attempt": attempt_dict
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get practice attempt detail error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch practice attempt"
        )


@router.get("/documents/{document_id}/stats")
@limiter.limit("30/minute")
async def get_practice_set_stats(
    request: Request,
    document_id: str,
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database)
):
    """
    Get aggregated statistics for a practice set.
    
    Returns:
    - Per-question statistics (attempts, last correct, mastery status)
    - Summary (total questions, mastered, needs practice)
    """
    try:
        user_id = current_user.get("user_id") or current_user.get("student_id") or current_user.get("id")
        
        # Detect if user is B2C
        user_type = current_user.get("user_type", "")
        is_b2c = current_user.get("is_b2c", False) or user_type == "b2c_user"
        
        # Aggregation pipeline
        pipeline = [
            {"$match": {
                "student_id": str(user_id),
                "document_id": document_id
            }},
            {"$sort": {"created_at": -1}},
            {"$group": {
                "_id": "$question_id",
                "attempts": {"$sum": 1},
                "correct_count": {"$sum": {"$cond": ["$is_correct", 1, 0]}},
                "last_attempt": {"$first": "$created_at"},
                "last_correct": {"$first": "$is_correct"},
                "last_answer": {"$first": "$student_answer"},
                "question_text": {"$first": "$question_text"},
                "difficulty": {"$first": "$difficulty"},
                "avg_time": {"$avg": {"$ifNull": ["$time_spent", 0]}}
            }}
        ]
        
        if is_b2c:
            stats = await db.b2c_aggregate("practice_attempts", pipeline)
        else:
            stats = await db.mongo_aggregate("practice_attempts", pipeline)
        
        stats_list = list(stats)
        
        # Convert for JSON serialization
        question_stats = []
        for s in stats_list:
            stat_dict = {
                "question_id": s.get("_id"),
                "attempts": s.get("attempts", 0),
                "correct_count": s.get("correct_count", 0),
                "last_attempt": s.get("last_attempt").isoformat() if s.get("last_attempt") else None,
                "last_correct": s.get("last_correct", False),
                "last_answer": s.get("last_answer"),
                "question_text": s.get("question_text", "")[:200] if s.get("question_text") else "",
                "difficulty": s.get("difficulty"),
                "avg_time": s.get("avg_time"),
                "mastered": s.get("last_correct", False) and s.get("correct_count", 0) >= 1
            }
            question_stats.append(stat_dict)
        
        # Summary statistics
        total_questions = len(question_stats)
        mastered = sum(1 for q in question_stats if q.get("mastered"))
        needs_practice = total_questions - mastered
        total_attempts = sum(q.get("attempts", 0) for q in question_stats)
        total_correct = sum(q.get("correct_count", 0) for q in question_stats)
        
        return {
            "success": True,
            "document_id": document_id,
            "question_stats": question_stats,
            "summary": {
                "total_questions": total_questions,
                "mastered": mastered,
                "needs_practice": needs_practice,
                "total_attempts": total_attempts,
                "total_correct": total_correct,
                "overall_accuracy": (total_correct / total_attempts * 100) if total_attempts > 0 else 0.0
            }
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get practice set stats error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch practice set statistics"
        )


@router.get("/teacher-feedback")
@limiter.limit("60/minute")
async def get_teacher_feedback_for_document(
    request: Request,
    document_id: str = Query(..., description="Practice set document ID"),
    current_user: Dict[str, Any] = Depends(get_current_user),
    db: DatabaseManager = Depends(get_database),
):
    """
    Get all teacher feedback for the current student's attempts on a document.
    Returns a map of question_id -> teacher_feedback so the student can view
    feedback from their teacher on each question.
    """
    try:
        user_id = current_user.get("user_id") or current_user.get("student_id") or current_user.get("id")

        # Detect if user is B2C
        user_type = current_user.get("user_type", "")
        is_b2c = current_user.get("is_b2c", False) or user_type == "b2c_user"

        filter_dict = {
            "student_id": str(user_id),
            "document_id": document_id,
            "teacher_feedback": {"$exists": True, "$ne": None},
        }

        if is_b2c:
            attempts = await db.b2c_find(
                "practice_attempts",
                filter_dict,
                sort=[("created_at", -1)],
            )
        else:
            attempts = await db.mongo_find(
                "practice_attempts",
                filter_dict,
                sort=[("created_at", -1)],
            )

        # Build a map keyed by question_id (keep only the latest per question)
        feedback_map: dict = {}
        for a in attempts:
            qid = a.get("question_id", "")
            if qid and qid not in feedback_map:
                tf = a.get("teacher_feedback", {})
                created_at = tf.get("created_at")
                updated_at = tf.get("updated_at")
                feedback_map[qid] = {
                    "question_id": qid,
                    "question_text": (a.get("question_text", "")[:200]
                                      if a.get("question_text") else ""),
                    "teacher_feedback": {
                        "text": tf.get("text", ""),
                        "tutor_name": tf.get("tutor_name", ""),
                        "created_at": created_at.isoformat()
                        if hasattr(created_at, "isoformat") else str(created_at or ""),
                        "updated_at": updated_at.isoformat()
                        if hasattr(updated_at, "isoformat") else str(updated_at or ""),
                    },
                }

        return {
            "success": True,
            "document_id": document_id,
            "feedback": list(feedback_map.values()),
            "total": len(feedback_map),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Get teacher feedback error: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to fetch teacher feedback",
        )
