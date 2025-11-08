"""
Document Processing Service - Handle PDF, Word, and Image files
Extracts text and prepares documents for RAG system
"""

import logging
import io
import base64
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import asyncio
from datetime import datetime

# Document loaders
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    convert_from_bytes = None
    PDF2IMAGE_AVAILABLE = False

try:
    import pypdfium2 as pdfium
    PDFIUM_AVAILABLE = True
except ImportError:
    pdfium = None
    PDFIUM_AVAILABLE = False

# LangChain text splitters (using new import path for LangChain 1.x)
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Robust document processing service for multiple file types
    Handles PDF, Word documents, and images with OCR
    """

    # Supported file types
    SUPPORTED_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc',
        'image/jpeg': 'image',
        'image/png': 'image',
        'image/jpg': 'image',
        'image/webp': 'image'
    }

    # Text splitting configuration
    CHUNK_SIZE = 1000
    CHUNK_OVERLAP = 200

    def __init__(self):
        """Initialize document processor"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.CHUNK_SIZE,
            chunk_overlap=self.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len
        )
        logger.info("📄 Document Processor initialized")

    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Process uploaded file and extract text

        Args:
            file_content: Binary file content
            filename: Original filename
            mime_type: MIME type of the file

        Returns:
            Dict with success status, text content, chunks, and metadata
        """
        try:
            # Validate file type
            if mime_type not in self.SUPPORTED_TYPES:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {mime_type}",
                    "supported_types": list(self.SUPPORTED_TYPES.keys())
                }

            file_type = self.SUPPORTED_TYPES[mime_type]

            logger.info(f"📥 Processing {file_type} file: {filename} ({len(file_content)} bytes)")

            # Route to appropriate processor
            if file_type == 'pdf':
                result = await self._process_pdf(file_content, filename)
            elif file_type in ['docx', 'doc']:
                result = await self._process_word(file_content, filename)
            elif file_type == 'image':
                result = await self._process_image(file_content, filename)
            else:
                return {
                    "success": False,
                    "error": f"Processor not implemented for {file_type}"
                }

            if not result["success"]:
                return result

            # Split text into chunks for RAG
            text = result["text"]
            chunks = self.text_splitter.split_text(text)

            logger.info(f"✅ Processed {filename}: {len(text)} chars -> {len(chunks)} chunks")

            return {
                "success": True,
                "filename": filename,
                "file_type": file_type,
                "text": text,
                "chunks": chunks,
                "num_chunks": len(chunks),
                "char_count": len(text),
                "metadata": {
                    "filename": filename,
                    "file_type": file_type,
                    "processed_at": datetime.now().isoformat(),
                    "chunk_size": self.CHUNK_SIZE,
                    "chunk_overlap": self.CHUNK_OVERLAP
                }
            }

        except Exception as e:
            logger.error(f"❌ Error processing file {filename}: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": f"Failed to process file: {str(e)}",
                "filename": filename
            }

    async def _process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from PDF file with enhanced extraction

        Args:
            content: PDF file bytes
            filename: Original filename

        Returns:
            Dict with success status and extracted text
        """
        try:
            if PdfReader is None:
                return {
                    "success": False,
                    "error": "PDF processing library not available. Install pypdf."
                }

            # Read PDF
            pdf_file = io.BytesIO(content)
            pdf_reader = PdfReader(pdf_file)

            # Extract text from all pages with multiple extraction methods
            text_parts = []
            total_chars = 0

            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    # Method 1: Standard extraction
                    page_text = page.extract_text()

                    # Method 2: Try with different extraction modes if available
                    if hasattr(page, 'extract_text'):
                        try:
                            # Try extraction with layout mode
                            alt_text = page.extract_text(extraction_mode="layout")
                            if len(alt_text) > len(page_text):
                                page_text = alt_text
                        except:
                            pass

                    # Clean and validate text
                    if page_text and page_text.strip():
                        # Remove excessive whitespace
                        page_text = '\n'.join(line.strip() for line in page_text.split('\n') if line.strip())

                        if len(page_text) > 50:  # Only add if substantial content
                            text_parts.append(f"[Page {page_num}]\n{page_text}")
                            total_chars += len(page_text)
                            logger.debug(f"📄 Page {page_num}: Extracted {len(page_text)} chars")
                        else:
                            logger.warning(f"⚠️ Page {page_num}: Only {len(page_text)} chars (might be image-based)")

                except Exception as e:
                    logger.warning(f"⚠️ Failed to extract text from page {page_num}: {str(e)}")
                    continue

            # Check if we got meaningful content
            if not text_parts:
                return {
                    "success": False,
                    "error": "❌ No text could be extracted from this PDF.\n\n📸 **This appears to be a scanned/image-based PDF.**\n\n**Solution:** Please convert the PDF pages to images (JPG/PNG) and upload them instead. I can read and analyze images directly using vision AI!"
                }

            if total_chars < 200:
                logger.warning(f"⚠️ Only extracted {total_chars} chars from PDF - likely image-based, will convert to images")

                # Try to convert PDF to images automatically
                images = []

                if PDF2IMAGE_AVAILABLE:
                    try:
                        logger.info("📸 Attempting automatic PDF-to-image conversion for all pages...")
                        images = convert_from_bytes(content, dpi=150)
                        if images:
                            logger.info(f"✅ Converted PDF to {len(images)} images via pdf2image")
                    except Exception as e:
                        logger.error(f"❌ PDF-to-image conversion with pdf2image failed: {str(e)}")
                        images = []

                if not images and PDFIUM_AVAILABLE:
                    try:
                        logger.info("📸 Falling back to pypdfium2 for PDF-to-image conversion...")
                        pdf_io = io.BytesIO(content)
                        pdf_document = pdfium.PdfDocument(pdf_io)
                        page_total = len(pdf_document)

                        for page_index in range(page_total):
                            page = pdf_document[page_index]
                            bitmap = page.render(scale=150 / 72)
                            pil_image = bitmap.to_pil()
                            bitmap.close()

                            if hasattr(page, 'close'):
                                page.close()

                            images.append(pil_image)

                        if hasattr(pdf_document, 'close'):
                            pdf_document.close()

                        if images:
                            logger.info(f"✅ Converted PDF to {len(images)} images via pypdfium2")
                    except Exception as e:
                        logger.error(f"❌ pypdfium2 conversion failed: {str(e)}")
                        images = []

                if images:
                    # Convert images to base64 for Vision API processing
                    image_texts = []
                    for i, img in enumerate(images, 1):
                        img_buffer = io.BytesIO()
                        img.save(img_buffer, format='PNG')
                        img_buffer.seek(0)

                        # Convert to base64
                        img_base64 = base64.b64encode(img_buffer.getvalue()).decode('utf-8')
                        img_data_url = f"data:image/png;base64,{img_base64}"

                        # Create a placeholder text indicating this page needs Vision API
                        image_texts.append(
                            f"[Page {i} - Image Content]\n"
                            f"{{IMAGE_DATA:{img_data_url}}}\n"
                            f"Note: This page contains visual content that will be analyzed using Vision AI.\n"
                        )

                    # Return success with image markers
                    combined_text = "\n\n".join(image_texts)
                    logger.info(f"📸 PDF converted to {len(images)} images for Vision AI processing")

                    return {
                        "success": True,
                        "text": combined_text,
                        "num_pages": len(pdf_reader.pages),
                        "extracted_pages": len(images),
                        "is_image_based": True,
                        "num_images": len(images)
                    }

                # If conversion failed or not available, return helpful error
                return {
                    "success": False,
                    "error": f"❌ Minimal text extracted ({total_chars} chars).\n\n📄 **Automatic PDF-to-image conversion was attempted but failed.**\n\n**Solution:** Please:\n1. Convert PDF pages to images (JPG/PNG)\n2. Upload the images instead\n3. I can read and analyze images directly!"
                }

            full_text = "\n\n".join(text_parts)
            logger.info(f"📄 Extracted {len(full_text)} characters from {len(text_parts)}/{len(pdf_reader.pages)} pages")

            return {
                "success": True,
                "text": full_text,
                "num_pages": len(pdf_reader.pages),
                "extracted_pages": len(text_parts)
            }

        except Exception as e:
            logger.error(f"❌ PDF processing error: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to read PDF: {str(e)}"
            }

    async def _process_word(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from Word document (.docx)

        Args:
            content: Word file bytes
            filename: Original filename

        Returns:
            Dict with success status and extracted text
        """
        try:
            if DocxDocument is None:
                return {
                    "success": False,
                    "error": "Word processing library not available. Install python-docx."
                }

            # Read Word document
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            if not text_parts:
                return {
                    "success": False,
                    "error": "No text could be extracted from Word document"
                }

            full_text = "\n\n".join(text_parts)
            logger.debug(f"📄 Extracted {len(full_text)} characters from Word document")

            return {
                "success": True,
                "text": full_text,
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables)
            }

        except Exception as e:
            logger.error(f"❌ Word processing error: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to read Word document: {str(e)}"
            }

    async def _process_image(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process image file - return base64 for vision API
        Note: Actual OCR can be added with pytesseract if needed

        Args:
            content: Image file bytes
            filename: Original filename

        Returns:
            Dict with success status and base64 image data
        """
        try:
            if Image is None:
                return {
                    "success": False,
                    "error": "Image processing library not available. Install Pillow."
                }

            # Validate image
            image = Image.open(io.BytesIO(content))
            width, height = image.size

            # Convert to base64 for vision API
            base64_data = base64.b64encode(content).decode('utf-8')
            mime_type = f"data:image/{image.format.lower()};base64,{base64_data}"

            logger.debug(f"🖼️ Processed image: {width}x{height}, format: {image.format}")

            # Return image data - the AI vision model will process it
            return {
                "success": True,
                "text": f"[Image: {filename} ({width}x{height})]",
                "image_data": mime_type,
                "width": width,
                "height": height,
                "format": image.format
            }

        except Exception as e:
            logger.error(f"❌ Image processing error: {str(e)}")
            return {
                "success": False,
                "error": f"Failed to process image: {str(e)}"
            }

    def validate_file(
        self,
        filename: str,
        mime_type: str,
        file_size: int,
        max_size_mb: int = 10
    ) -> Dict[str, Any]:
        """
        Validate file before processing

        Args:
            filename: Original filename
            mime_type: MIME type
            file_size: File size in bytes
            max_size_mb: Maximum allowed size in MB

        Returns:
            Dict with validation result
        """
        try:
            # Check file type
            if mime_type not in self.SUPPORTED_TYPES:
                return {
                    "valid": False,
                    "error": f"Unsupported file type: {mime_type}",
                    "supported_types": list(self.SUPPORTED_TYPES.keys())
                }

            # Check file size
            max_size_bytes = max_size_mb * 1024 * 1024
            if file_size > max_size_bytes:
                return {
                    "valid": False,
                    "error": f"File too large: {file_size / (1024*1024):.2f}MB (max {max_size_mb}MB)"
                }

            # Check filename
            if not filename or len(filename) > 255:
                return {
                    "valid": False,
                    "error": "Invalid filename"
                }

            return {
                "valid": True,
                "file_type": self.SUPPORTED_TYPES[mime_type]
            }

        except Exception as e:
            logger.error(f"❌ Validation error: {str(e)}")
            return {
                "valid": False,
                "error": f"Validation failed: {str(e)}"
            }


# Global instance
_document_processor = None


def get_document_processor() -> DocumentProcessor:
    """Get or create DocumentProcessor instance (singleton)"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
